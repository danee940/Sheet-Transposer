"""Tests for the Flask web layer in app.py."""

from io import BytesIO
from urllib.parse import unquote

import pytest
from docx import Document

import app as app_module
from app import app

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def client():
    app.config["TESTING"] = True
    with app.test_client() as test_client:
        yield test_client


def _chord_docx_bytes():
    document = Document()
    document.add_paragraph("C G Am F")
    document.add_paragraph("These are the lyrics")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _upload(data_bytes, filename="song.docx"):
    return {"file": (BytesIO(data_bytes), filename)}


def test_index_lists_keys(client):
    response = client.get("/")
    assert response.status_code == 200
    assert b"Chord Transposer" in response.data


def test_robots_txt(client):
    response = client.get("/robots.txt")
    assert response.status_code == 200
    assert response.mimetype == "text/plain"
    assert "Sitemap:" in response.get_data(as_text=True)


def test_sitemap_xml(client):
    response = client.get("/sitemap.xml")
    assert response.status_code == 200
    assert response.mimetype == "application/xml"
    assert "<urlset" in response.get_data(as_text=True)


def test_transpose_missing_file(client):
    response = client.post(
        "/transpose",
        data={"current_key": "C", "target_key": "D"},
        content_type="multipart/form-data",
    )
    assert response.status_code == 400
    assert "choose a .docx" in response.get_json()["error"]


def test_transpose_wrong_extension(client):
    data = _upload(b"not a docx", filename="song.txt")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Only .docx" in response.get_json()["error"]


def test_transpose_missing_keys(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "", "target_key": ""})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "keys are required" in response.get_json()["error"]


def test_transpose_unsupported_format(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D", "format": "txt"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Unsupported output format" in response.get_json()["error"]


def test_transpose_empty_file(client):
    data = _upload(b"", filename="empty.docx")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "empty" in response.get_json()["error"]


def test_transpose_invalid_key(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "X", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "not a valid key" in response.get_json()["error"]


def test_transpose_invalid_docx(client):
    data = _upload(b"this is not a real docx", filename="bad.docx")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "valid .docx" in response.get_json()["error"]


def test_transpose_docx_success(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 200
    assert response.mimetype == DOCX_MIMETYPE
    assert "song_D.docx" in response.headers["Content-Disposition"]
    assert unquote(response.headers["X-Transpose-From"]) == "C"
    assert unquote(response.headers["X-Transpose-To"]) == "D"
    assert "C->D" in unquote(response.headers["X-Transpose-Changes"])


def test_transpose_pdf_success(client, monkeypatch):
    monkeypatch.setattr(app_module, "convert_docx_to_pdf", lambda docx_bytes: b"%PDF-1.4 fake")

    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D", "format": "pdf"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 200
    assert response.mimetype == "application/pdf"
    assert "song_D.pdf" in response.headers["Content-Disposition"]
    assert response.data == b"%PDF-1.4 fake"


def test_transpose_pdf_conversion_failure(client, monkeypatch):
    def _raise(_docx_bytes):
        raise app_module.PdfConversionError("PDF conversion timed out.")

    monkeypatch.setattr(app_module, "convert_docx_to_pdf", _raise)

    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D", "format": "pdf"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 503
    assert "timed out" in response.get_json()["error"]
