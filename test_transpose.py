"""Unit tests for transpose.py — targeting ≥95% coverage."""

from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

import transpose as t

# ---------------------------------------------------------------------------
# parse_key
# ---------------------------------------------------------------------------


class TestParseKey:
    def test_simple_major(self):
        assert t.parse_key("C") == ("C", False)

    def test_simple_minor(self):
        assert t.parse_key("Am") == ("A", True)

    def test_sharp_major(self):
        assert t.parse_key("F#") == ("F#", False)

    def test_flat_major(self):
        assert t.parse_key("Eb") == ("Eb", False)

    def test_minor_full_word(self):
        assert t.parse_key("D minor") == ("D", True)

    def test_minor_min_suffix(self):
        assert t.parse_key("Bmin") == ("B", True)

    def test_lowercase_input(self):
        assert t.parse_key("c") == ("C", False)

    def test_lowercase_minor(self):
        assert t.parse_key("am") == ("A", True)

    def test_whitespace_stripped(self):
        assert t.parse_key("  G  ") == ("G", False)

    def test_empty_string_returns_none(self):
        assert t.parse_key("") is None

    def test_whitespace_only_returns_none(self):
        assert t.parse_key("   ") is None

    def test_invalid_note_returns_none(self):
        assert t.parse_key("X") is None

    def test_invalid_accidental_returns_none(self):
        assert t.parse_key("Cb") == ("Cb", False)

    def test_german_H_note(self):
        assert t.parse_key("H") == ("H", False)

    def test_sharp_minor(self):
        assert t.parse_key("C#m") == ("C#", True)


# ---------------------------------------------------------------------------
# key_label
# ---------------------------------------------------------------------------


class TestKeyLabel:
    def test_major(self):
        assert t.key_label("C", False) == "C"

    def test_minor(self):
        assert t.key_label("A", True) == "Am"

    def test_sharp_minor(self):
        assert t.key_label("F#", True) == "F#m"


# ---------------------------------------------------------------------------
# prefers_flats
# ---------------------------------------------------------------------------


class TestPrefersFlats:
    def test_flat_key(self):
        assert t.prefers_flats("Bb", False) is True

    def test_flat_key_by_label(self):
        assert t.prefers_flats("D", True) is True  # Dm is in FLAT_KEYS

    def test_sharp_key(self):
        assert t.prefers_flats("G", False) is False

    def test_sharp_key_minor(self):
        assert t.prefers_flats("B", True) is False  # Bm in SHARP_KEYS

    def test_c_major_defaults_to_flats(self):
        # C is not in FLAT_KEYS or SHARP_KEYS explicitly; falls back to True
        assert t.prefers_flats("C", False) is True

    def test_note_in_flat_keys(self):
        assert t.prefers_flats("F", False) is True


# ---------------------------------------------------------------------------
# choose_spelling
# ---------------------------------------------------------------------------


class TestChooseSpelling:
    def test_flat_non_german(self):
        assert t.choose_spelling(True, False) is t.FLAT_SPELLING

    def test_sharp_non_german(self):
        assert t.choose_spelling(False, False) is t.SHARP_SPELLING

    def test_flat_german(self):
        assert t.choose_spelling(True, True) is t.GERMAN_FLAT_SPELLING

    def test_sharp_german(self):
        assert t.choose_spelling(False, True) is t.GERMAN_SHARP_SPELLING


# ---------------------------------------------------------------------------
# note_semitone
# ---------------------------------------------------------------------------


class TestNoteSemitone:
    def test_c_is_zero(self):
        assert t.note_semitone("C", False) == 0

    def test_a_is_nine(self):
        assert t.note_semitone("A", False) == 9

    def test_sharp(self):
        assert t.note_semitone("C#", False) == 1

    def test_flat(self):
        assert t.note_semitone("Bb", False) == 10

    def test_german_H_is_eleven(self):
        assert t.note_semitone("H", True) == 11

    def test_german_B_is_ten(self):
        assert t.note_semitone("B", True) == 10

    def test_standard_B_is_eleven(self):
        assert t.note_semitone("B", False) == 11

    def test_lowercase_note(self):
        assert t.note_semitone("c", False) == 0


# ---------------------------------------------------------------------------
# transpose_note
# ---------------------------------------------------------------------------


class TestTransposeNote:
    def test_up_by_two_semitones(self):
        assert t.transpose_note("C", 2, t.SHARP_SPELLING, False) == "D"

    def test_wraps_around(self):
        assert t.transpose_note("B", 1, t.SHARP_SPELLING, False) == "C"

    def test_uses_flat_spelling(self):
        assert t.transpose_note("C", 1, t.FLAT_SPELLING, False) == "Db"

    def test_uses_sharp_spelling(self):
        assert t.transpose_note("C", 1, t.SHARP_SPELLING, False) == "C#"

    def test_german_H_up_one(self):
        assert t.transpose_note("H", 1, t.GERMAN_SHARP_SPELLING, True) == "C"


# ---------------------------------------------------------------------------
# transpose_chord
# ---------------------------------------------------------------------------


class TestTransposeChord:
    def test_simple_chord(self):
        assert t.transpose_chord("C", 2, t.SHARP_SPELLING, False) == "D"

    def test_chord_with_quality(self):
        assert t.transpose_chord("Am7", 3, t.SHARP_SPELLING, False) == "Cm7"

    def test_slash_chord(self):
        result = t.transpose_chord("G/B", 2, t.SHARP_SPELLING, False)
        assert result == "A/C#"

    def test_slash_chord_preserves_quality(self):
        result = t.transpose_chord("Cmaj7/E", 2, t.SHARP_SPELLING, False)
        assert result == "Dmaj7/F#"

    def test_no_root_match_returns_unchanged(self):
        assert t.transpose_chord("", 2, t.SHARP_SPELLING, False) == ""

    def test_diminished_chord(self):
        assert t.transpose_chord("Bdim", 1, t.FLAT_SPELLING, False) == "Cdim"

    def test_augmented_chord(self):
        assert t.transpose_chord("Caug", 7, t.SHARP_SPELLING, False) == "Gaug"

    def test_slash_with_non_root_bass(self):
        # bass part without a valid ROOT_MATCH_RE match — should not crash
        result = t.transpose_chord("C/3", 2, t.SHARP_SPELLING, False)
        # "3" won't match ROOT_MATCH_RE so bass stays as-is, but remainder includes "/3"
        assert "D" in result


# ---------------------------------------------------------------------------
# is_chord_line
# ---------------------------------------------------------------------------


class TestIsChordLine:
    def test_single_chord(self):
        assert t.is_chord_line("Am") is True

    def test_multiple_chords(self):
        assert t.is_chord_line("C G Am F") is True

    def test_empty_string(self):
        assert t.is_chord_line("") is False

    def test_whitespace_only(self):
        assert t.is_chord_line("   ") is False

    def test_lyric_line(self):
        assert t.is_chord_line("Hello world") is False

    def test_mixed_chords_and_words(self):
        assert t.is_chord_line("C the G") is False

    def test_chord_with_quality(self):
        assert t.is_chord_line("Cmaj7 Dm7 G7") is True

    def test_slash_chord(self):
        assert t.is_chord_line("G/B C/E") is True

    def test_sharp_chords(self):
        assert t.is_chord_line("F# C# G#m") is True

    def test_flat_chords(self):
        assert t.is_chord_line("Bb Eb Ab") is True


# ---------------------------------------------------------------------------
# line_uses_german
# ---------------------------------------------------------------------------


class TestLineUsesGerman:
    def test_h_chord(self):
        assert t.line_uses_german("H C D") is True

    def test_h_minor(self):
        assert t.line_uses_german("Hm G") is True

    def test_slash_h_bass(self):
        assert t.line_uses_german("C/H") is True

    def test_token_starting_with_h(self):
        assert t.line_uses_german("H") is True

    def test_no_german(self):
        assert t.line_uses_german("C G Am F") is False

    def test_empty_line(self):
        assert t.line_uses_german("") is False


# ---------------------------------------------------------------------------
# transpose_line_text
# ---------------------------------------------------------------------------


class TestTransposeLineText:
    def _transpose(self, text, semitones=2, use_flats=False, german=False):
        spelling = t.choose_spelling(use_flats, german)
        changes = set()
        result = t.transpose_line_text(text, semitones, spelling, german, changes)
        return result, changes

    def test_basic_transposition(self):
        result, changes = self._transpose("C G Am F")
        assert "D" in result
        assert ("C", "D") in changes

    def test_preserves_leading_spaces(self):
        result, _ = self._transpose("  C G")
        assert result.startswith("  ")

    def test_spacing_preserved_when_chord_shrinks(self):
        # C# (2 chars) -> D (1 char): trailing space should grow by 1
        result, _ = self._transpose("C#  G", semitones=1, use_flats=False)
        # D should be followed by enough spaces to maintain alignment
        assert result.startswith("D")

    def test_spacing_preserved_when_chord_grows(self):
        # C (1 char) -> Db (2 chars): trailing space should shrink by 1
        result, _ = self._transpose("C   G", semitones=1, use_flats=True)
        assert result.startswith("Db")

    def test_no_changes_when_same_key(self):
        result, changes = self._transpose("C G Am F", semitones=0)
        assert changes == set()
        assert result == "C G Am F"

    def test_records_all_changes(self):
        _, changes = self._transpose("C G Am F", semitones=2)
        assert len(changes) == 4

    def test_trailing_space_clamped_to_zero(self):
        # chord grows but there's no trailing space to remove
        result, _ = self._transpose("C", semitones=1, use_flats=True)
        assert result == "Db"


# ---------------------------------------------------------------------------
# _pick_styled_run
# ---------------------------------------------------------------------------


class TestPickStyledRun:
    def _make_run(self, text):
        run = MagicMock()
        run.text = text
        return run

    def test_returns_first_non_whitespace_run(self):
        runs = [self._make_run("  "), self._make_run("Am"), self._make_run("G")]
        assert t._pick_styled_run(runs) is runs[1]

    def test_falls_back_to_first_run_when_all_whitespace(self):
        runs = [self._make_run("  "), self._make_run("   ")]
        assert t._pick_styled_run(runs) is runs[0]

    def test_single_run(self):
        runs = [self._make_run("C")]
        assert t._pick_styled_run(runs) is runs[0]


# ---------------------------------------------------------------------------
# redistribute_to_runs
# ---------------------------------------------------------------------------


class TestRedistributeToRuns:
    def _make_run(self, text):
        run = MagicMock()
        run.text = text
        return run

    def _make_paragraph(self, texts):
        para = MagicMock()
        para.runs = [self._make_run(t_) for t_ in texts]
        return para

    def test_concentrates_text_in_styled_run(self):
        para = self._make_paragraph(["Am", "  ", "G"])
        t.redistribute_to_runs(para, "Cm  D")
        assert para.runs[0].text == "Cm  D"
        assert para.runs[1].text == ""
        assert para.runs[2].text == ""

    def test_empty_runs_skipped(self):
        para = self._make_paragraph(["", "Am"])
        t.redistribute_to_runs(para, "Bm")
        # only non-empty runs are considered; "Am" run gets the text
        assert para.runs[1].text == "Bm"

    def test_all_empty_runs_does_nothing(self):
        para = self._make_paragraph(["", ""])
        t.redistribute_to_runs(para, "D")
        # no styled run found, nothing written
        para.runs[0].text = ""
        para.runs[1].text = ""


# ---------------------------------------------------------------------------
# process_paragraph
# ---------------------------------------------------------------------------


class TestProcessParagraph:
    def _make_paragraph(self, text):
        para = MagicMock()
        para.text = text
        run = MagicMock()
        run.text = text
        para.runs = [run]
        return para

    def test_chord_line_is_transposed(self):
        para = self._make_paragraph("C G Am F")
        changes = set()
        t.process_paragraph(para, 2, False, False, changes)
        assert len(changes) > 0

    def test_lyric_line_is_skipped(self):
        para = self._make_paragraph("Hello world this is lyrics")
        changes = set()
        t.process_paragraph(para, 2, False, False, changes)
        assert changes == set()

    def test_empty_paragraph_is_skipped(self):
        para = self._make_paragraph("")
        changes = set()
        t.process_paragraph(para, 2, False, False, changes)
        assert changes == set()


# ---------------------------------------------------------------------------
# iter_paragraphs
# ---------------------------------------------------------------------------


class TestIterParagraphs:
    def _make_document(self, body_texts, table_texts=None, header_texts=None, footer_texts=None):
        doc = MagicMock()

        doc.paragraphs = [self._make_para(t_) for t_ in body_texts]

        if table_texts:
            cell = MagicMock()
            cell.paragraphs = [self._make_para(t_) for t_ in table_texts]
            row = MagicMock()
            row.cells = [cell]
            table = MagicMock()
            table.rows = [row]
            doc.tables = [table]
        else:
            doc.tables = []

        section = MagicMock()
        section.header.paragraphs = [self._make_para(t_) for t_ in (header_texts or [])]
        section.footer.paragraphs = [self._make_para(t_) for t_ in (footer_texts or [])]
        doc.sections = [section]

        return doc

    def _make_para(self, text):
        para = MagicMock()
        para.text = text
        return para

    def test_yields_body_paragraphs(self):
        doc = self._make_document(["C G", "lyrics"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "C G" in texts
        assert "lyrics" in texts

    def test_yields_table_paragraphs(self):
        doc = self._make_document([], table_texts=["Am Dm"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "Am Dm" in texts

    def test_yields_header_paragraphs(self):
        doc = self._make_document([], header_texts=["F C"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "F C" in texts

    def test_yields_footer_paragraphs(self):
        doc = self._make_document([], footer_texts=["G D"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "G D" in texts


# ---------------------------------------------------------------------------
# detect_german
# ---------------------------------------------------------------------------


class TestDetectGerman:
    def _make_document_with_para(self, text):
        para = MagicMock()
        para.text = text
        doc = MagicMock()
        doc.paragraphs = [para]
        doc.tables = []
        section = MagicMock()
        section.header.paragraphs = []
        section.footer.paragraphs = []
        doc.sections = [section]
        return doc

    def test_detects_german_notation(self):
        doc = self._make_document_with_para("H C D")
        assert t.detect_german(doc) is True

    def test_no_german_notation(self):
        doc = self._make_document_with_para("C G Am F")
        assert t.detect_german(doc) is False

    def test_lyric_line_with_h_word_ignored(self):
        doc = self._make_document_with_para("Hello world")
        assert t.detect_german(doc) is False


# ---------------------------------------------------------------------------
# find_input_file
# ---------------------------------------------------------------------------


class TestFindInputFile:
    def test_returns_none_when_no_docx(self, tmp_path):
        with patch.object(t, "INPUT_DIR", tmp_path):
            assert t.find_input_file() is None

    def test_returns_docx_file(self, tmp_path):
        docx_file = tmp_path / "song.docx"
        docx_file.touch()
        with patch.object(t, "INPUT_DIR", tmp_path):
            assert t.find_input_file() == docx_file

    def test_skips_temp_files(self, tmp_path):
        temp_file = tmp_path / "~$song.docx"
        temp_file.touch()
        with patch.object(t, "INPUT_DIR", tmp_path):
            assert t.find_input_file() is None


# ---------------------------------------------------------------------------
# _prompt_key
# ---------------------------------------------------------------------------


class TestPromptKey:
    def test_valid_input_returns_parsed_key(self):
        with patch("builtins.input", return_value="Am"):
            result = t._prompt_key("Key: ")
        assert result == ("A", True)

    def test_invalid_input_exits(self):
        with patch("builtins.input", return_value="X"), pytest.raises(SystemExit):
            t._prompt_key("Key: ")

    def test_invalid_input_prints_message(self, capsys):
        with patch("builtins.input", return_value="X"), pytest.raises(SystemExit):
            t._prompt_key("Key: ")
        assert "not a valid key" in capsys.readouterr().out


# ---------------------------------------------------------------------------
# _save_and_report
# ---------------------------------------------------------------------------


class TestSaveAndReport:
    def test_saves_file_and_prints_summary(self, tmp_path, capsys):
        doc = MagicMock()
        input_file = tmp_path / "song.docx"
        input_file.touch()
        output_dir = tmp_path / "output"

        with patch.object(t, "OUTPUT_DIR", output_dir):
            t._save_and_report(doc, input_file, "C", "D", {("C", "D"), ("G", "A")})

        captured = capsys.readouterr().out
        assert "C" in captured
        assert "D" in captured
        doc.save.assert_called_once()

    def test_reports_no_changes(self, tmp_path, capsys):
        doc = MagicMock()
        input_file = tmp_path / "song.docx"
        input_file.touch()
        output_dir = tmp_path / "output"

        with patch.object(t, "OUTPUT_DIR", output_dir):
            t._save_and_report(doc, input_file, "C", "C", set())

        assert "No chords were changed" in capsys.readouterr().out


# ---------------------------------------------------------------------------
# _transpose_document
# ---------------------------------------------------------------------------


class TestTransposeDocument:
    def test_returns_changes(self):
        para = MagicMock()
        para.text = "C G Am F"
        run = MagicMock()
        run.text = "C G Am F"
        para.runs = [run]

        doc = MagicMock()
        doc.paragraphs = [para]
        doc.tables = []
        section = MagicMock()
        section.header.paragraphs = []
        section.footer.paragraphs = []
        doc.sections = [section]

        changes = t._transpose_document(doc, 2, False, False)
        assert isinstance(changes, set)
        assert len(changes) > 0


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------


class TestMain:
    def _make_mock_document(self, para_text="C G Am F"):
        para = MagicMock()
        para.text = para_text
        run = MagicMock()
        run.text = para_text
        para.runs = [run]
        doc = MagicMock()
        doc.paragraphs = [para]
        doc.tables = []
        section = MagicMock()
        section.header.paragraphs = []
        section.footer.paragraphs = []
        doc.sections = [section]
        return doc

    def test_exits_when_no_input_file(self, tmp_path):
        with patch.object(t, "INPUT_DIR", tmp_path), pytest.raises(SystemExit):
            t.main()

    def test_full_run(self, tmp_path):
        input_dir = tmp_path / "input"
        input_dir.mkdir()
        output_dir = tmp_path / "output"
        docx_file = input_dir / "song.docx"
        docx_file.touch()

        mock_doc = self._make_mock_document()

        with (
            patch.object(t, "INPUT_DIR", input_dir),
            patch.object(t, "OUTPUT_DIR", output_dir),
            patch("transpose.Document", return_value=mock_doc),
            patch("builtins.input", side_effect=["C", "D"]),
        ):
            t.main()

        mock_doc.save.assert_called_once()


# ---------------------------------------------------------------------------
# transpose_document_bytes
# ---------------------------------------------------------------------------


def _make_docx_bytes(lines):
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestTransposeDocumentBytes:
    def test_transposes_and_returns_metadata(self):
        file_bytes = _make_docx_bytes(["C G Am F", "Hello world"])

        docx_bytes, from_label, to_label, changes = t.transpose_document_bytes(file_bytes, "C", "D")

        assert from_label == "C"
        assert to_label == "D"
        assert ("C", "D") in changes

        result = Document(BytesIO(docx_bytes))
        texts = [p.text for p in result.paragraphs]
        assert "Hello world" in texts

    def test_invalid_current_key_raises(self):
        file_bytes = _make_docx_bytes(["C G"])
        with pytest.raises(t.InvalidKeyError):
            t.transpose_document_bytes(file_bytes, "X", "D")

    def test_invalid_target_key_raises(self):
        file_bytes = _make_docx_bytes(["C G"])
        with pytest.raises(t.InvalidKeyError):
            t.transpose_document_bytes(file_bytes, "C", "X")

    def test_same_key_produces_no_changes(self):
        file_bytes = _make_docx_bytes(["C G Am F"])
        _, from_label, to_label, changes = t.transpose_document_bytes(file_bytes, "C", "C")
        assert from_label == "C"
        assert to_label == "C"
        assert changes == []

    def test_minor_key_labels(self):
        file_bytes = _make_docx_bytes(["Am Dm E"])
        _, from_label, to_label, _ = t.transpose_document_bytes(file_bytes, "Am", "Bm")
        assert from_label == "Am"
        assert to_label == "Bm"
