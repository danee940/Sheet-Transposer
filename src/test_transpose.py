"""Unit tests for transpose.py — targeting ≥95% coverage."""
# pylint: disable=redefined-outer-name,missing-function-docstring

from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
import requests
from docx import Document

import transpose as t

# ---------------------------------------------------------------------------
# parse_key
# ---------------------------------------------------------------------------


class TestParseKey:
    def test_simple_major(self):
        assert t.parse_key("C") == ("C", False)

    def test_simple_minor(self):
        assert t.parse_key("Am") == ("A", True)

    def test_sharp_major(self):
        assert t.parse_key("F#") == ("F#", False)

    def test_flat_major(self):
        assert t.parse_key("Eb") == ("Eb", False)

    def test_minor_full_word(self):
        assert t.parse_key("D minor") == ("D", True)

    def test_minor_min_suffix(self):
        assert t.parse_key("Bmin") == ("B", True)

    def test_lowercase_input(self):
        assert t.parse_key("c") == ("C", False)

    def test_lowercase_minor(self):
        assert t.parse_key("am") == ("A", True)

    def test_whitespace_stripped(self):
        assert t.parse_key("  G  ") == ("G", False)

    def test_empty_string_returns_none(self):
        assert t.parse_key("") is None

    def test_whitespace_only_returns_none(self):
        assert t.parse_key("   ") is None

    def test_invalid_note_returns_none(self):
        assert t.parse_key("X") is None

    def test_invalid_accidental_returns_none(self):
        assert t.parse_key("Cb") == ("Cb", False)

    def test_german_H_note(self):
        assert t.parse_key("H") == ("H", False)

    def test_sharp_minor(self):
        assert t.parse_key("C#m") == ("C#", True)

    def test_german_es_note(self):
        assert t.parse_key("Es") == ("Es", False)

    def test_german_as_minor(self):
        assert t.parse_key("Asm") == ("As", True)

    def test_german_cis_note(self):
        assert t.parse_key("Cis") == ("Cis", False)

    def test_german_lowercase_es(self):
        assert t.parse_key("es") == ("Es", False)

    def test_german_name_with_trailing_word_falls_through(self):
        assert t.parse_key("Est") == ("E", False)


# ---------------------------------------------------------------------------
# key_label
# ---------------------------------------------------------------------------


class TestKeyLabel:
    def test_major(self):
        assert t.key_label("C", False) == "C"

    def test_minor(self):
        assert t.key_label("A", True) == "Am"

    def test_sharp_minor(self):
        assert t.key_label("F#", True) == "F#m"


# ---------------------------------------------------------------------------
# prefers_flats
# ---------------------------------------------------------------------------


class TestPrefersFlats:
    def test_flat_key(self):
        assert t.prefers_flats("Bb", False) is True

    def test_flat_key_by_label(self):
        assert t.prefers_flats("D", True) is True  # Dm is in FLAT_KEYS

    def test_sharp_key(self):
        assert t.prefers_flats("G", False) is False

    def test_sharp_key_minor(self):
        assert t.prefers_flats("B", True) is False  # Bm in SHARP_KEYS

    def test_c_major_defaults_to_flats(self):
        # C is not in FLAT_KEYS or SHARP_KEYS explicitly; falls back to True
        assert t.prefers_flats("C", False) is True

    def test_note_in_flat_keys(self):
        assert t.prefers_flats("F", False) is True


# ---------------------------------------------------------------------------
# choose_spelling
# ---------------------------------------------------------------------------


class TestChooseSpelling:
    def test_flat_non_german(self):
        assert t.choose_spelling(True, False) is t.FLAT_SPELLING

    def test_sharp_non_german(self):
        assert t.choose_spelling(False, False) is t.SHARP_SPELLING

    def test_flat_german(self):
        assert t.choose_spelling(True, True) is t.GERMAN_FLAT_SPELLING

    def test_sharp_german(self):
        assert t.choose_spelling(False, True) is t.GERMAN_SHARP_SPELLING


# ---------------------------------------------------------------------------
# key_semitone
# ---------------------------------------------------------------------------


class TestKeySemitone:
    def test_german_only_name_resolves_without_german_flag(self):
        assert t.key_semitone("Es", False) == 3

    def test_german_only_name_resolves_with_german_flag(self):
        assert t.key_semitone("Es", True) == 3

    def test_cis_resolves_without_german_flag(self):
        assert t.key_semitone("Cis", False) == 1

    def test_english_b_stays_natural(self):
        assert t.key_semitone("B", False) == 11

    def test_german_b_is_flat(self):
        assert t.key_semitone("B", True) == 10


# ---------------------------------------------------------------------------
# note_semitone
# ---------------------------------------------------------------------------


class TestNoteSemitone:
    def test_c_is_zero(self):
        assert t.note_semitone("C", False) == 0

    def test_a_is_nine(self):
        assert t.note_semitone("A", False) == 9

    def test_sharp(self):
        assert t.note_semitone("C#", False) == 1

    def test_flat(self):
        assert t.note_semitone("Bb", False) == 10

    def test_german_H_is_eleven(self):
        assert t.note_semitone("H", True) == 11

    def test_german_B_is_ten(self):
        assert t.note_semitone("B", True) == 10

    def test_standard_B_is_eleven(self):
        assert t.note_semitone("B", False) == 11

    def test_lowercase_note(self):
        assert t.note_semitone("c", False) == 0


# ---------------------------------------------------------------------------
# transpose_note
# ---------------------------------------------------------------------------


class TestTransposeNote:
    def test_up_by_two_semitones(self):
        assert t.transpose_note("C", 2, t.SHARP_SPELLING, False) == "D"

    def test_wraps_around(self):
        assert t.transpose_note("B", 1, t.SHARP_SPELLING, False) == "C"

    def test_uses_flat_spelling(self):
        assert t.transpose_note("C", 1, t.FLAT_SPELLING, False) == "Db"

    def test_uses_sharp_spelling(self):
        assert t.transpose_note("C", 1, t.SHARP_SPELLING, False) == "C#"

    def test_german_H_up_one(self):
        assert t.transpose_note("H", 1, t.GERMAN_SHARP_SPELLING, True) == "C"


# ---------------------------------------------------------------------------
# transpose_chord
# ---------------------------------------------------------------------------


class TestTransposeChord:
    def test_simple_chord(self):
        assert t.transpose_chord("C", 2, t.SHARP_SPELLING, False) == "D"

    def test_chord_with_quality(self):
        assert t.transpose_chord("Am7", 3, t.SHARP_SPELLING, False) == "Cm7"

    def test_slash_chord(self):
        result = t.transpose_chord("G/B", 2, t.SHARP_SPELLING, False)
        assert result == "A/C#"

    def test_slash_chord_preserves_quality(self):
        result = t.transpose_chord("Cmaj7/E", 2, t.SHARP_SPELLING, False)
        assert result == "Dmaj7/F#"

    def test_no_root_match_returns_unchanged(self):
        assert t.transpose_chord("", 2, t.SHARP_SPELLING, False) == ""

    def test_diminished_chord(self):
        assert t.transpose_chord("Bdim", 1, t.FLAT_SPELLING, False) == "Cdim"

    def test_augmented_chord(self):
        assert t.transpose_chord("Caug", 7, t.SHARP_SPELLING, False) == "Gaug"

    def test_slash_with_non_root_bass(self):
        # bass part without a valid ROOT_MATCH_RE match — should not crash
        result = t.transpose_chord("C/3", 2, t.SHARP_SPELLING, False)
        # "3" won't match ROOT_MATCH_RE so bass stays as-is, but remainder includes "/3"
        assert "D" in result


# ---------------------------------------------------------------------------
# is_chord_line
# ---------------------------------------------------------------------------


class TestIsChordLine:
    def test_single_chord(self):
        assert t.is_chord_line("Am") is True

    def test_multiple_chords(self):
        assert t.is_chord_line("C G Am F") is True

    def test_empty_string(self):
        assert t.is_chord_line("") is False

    def test_whitespace_only(self):
        assert t.is_chord_line("   ") is False

    def test_lyric_line(self):
        assert t.is_chord_line("Hello world") is False

    def test_mixed_chords_and_words(self):
        assert t.is_chord_line("C the G") is False

    def test_chord_with_quality(self):
        assert t.is_chord_line("Cmaj7 Dm7 G7") is True

    def test_slash_chord(self):
        assert t.is_chord_line("G/B C/E") is True

    def test_sharp_chords(self):
        assert t.is_chord_line("F# C# G#m") is True

    def test_flat_chords(self):
        assert t.is_chord_line("Bb Eb Ab") is True


# ---------------------------------------------------------------------------
# line_uses_german
# ---------------------------------------------------------------------------


class TestLineUsesGerman:
    def test_h_chord(self):
        assert t.line_uses_german("H C D") is True

    def test_h_minor(self):
        assert t.line_uses_german("Hm G") is True

    def test_slash_h_bass(self):
        assert t.line_uses_german("C/H") is True

    def test_token_starting_with_h(self):
        assert t.line_uses_german("H") is True

    def test_no_german(self):
        assert t.line_uses_german("C G Am F") is False

    def test_empty_line(self):
        assert t.line_uses_german("") is False


# ---------------------------------------------------------------------------
# transpose_line_text
# ---------------------------------------------------------------------------


class TestTransposeLineText:
    def _transpose(self, text, semitones=2, use_flats=False, german=False):
        spelling = t.choose_spelling(use_flats, german)
        changes = set()
        result = t.transpose_line_text(text, semitones, spelling, german, changes)
        return result, changes

    def test_basic_transposition(self):
        result, changes = self._transpose("C G Am F")
        assert "D" in result
        assert ("C", "D") in changes

    def test_preserves_leading_spaces(self):
        result, _ = self._transpose("  C G")
        assert result.startswith("  ")

    def test_spacing_preserved_when_chord_shrinks(self):
        # C# (2 chars) -> D (1 char): trailing space should grow by 1
        result, _ = self._transpose("C#  G", semitones=1, use_flats=False)
        # D should be followed by enough spaces to maintain alignment
        assert result.startswith("D")

    def test_spacing_preserved_when_chord_grows(self):
        # C (1 char) -> Db (2 chars): trailing space should shrink by 1
        result, _ = self._transpose("C   G", semitones=1, use_flats=True)
        assert result.startswith("Db")

    def test_no_changes_when_same_key(self):
        result, changes = self._transpose("C G Am F", semitones=0)
        assert changes == set()
        assert result == "C G Am F"

    def test_records_all_changes(self):
        _, changes = self._transpose("C G Am F", semitones=2)
        assert len(changes) == 4

    def test_trailing_space_clamped_to_zero(self):
        # chord grows but there's no trailing space to remove
        result, _ = self._transpose("C", semitones=1, use_flats=True)
        assert result == "Db"


# ---------------------------------------------------------------------------
# transpose_line_with_owners
# ---------------------------------------------------------------------------


class TestTransposeLineWithOwners:
    def _run(self, text, owners, semitones=2, use_flats=False, german=False):
        spelling = t.choose_spelling(use_flats, german)
        return t.transpose_line_with_owners(text, owners, semitones, spelling, german)

    def test_owners_length_matches_text(self):
        text = "C G Am"
        owners = [0] * len(text)
        new_text, new_owners = self._run(text, owners)
        assert len(new_text) == len(new_owners)

    def test_leading_spaces_keep_owner(self):
        text = "  C"
        owners = [0, 0, 1]
        new_text, new_owners = self._run(text, owners)
        assert new_text.startswith("  ")
        assert new_owners[0] == 0

    def test_root_owner_taken_from_first_char(self):
        text = "Dm7"
        owners = [0, 0, 1]
        new_text, new_owners = self._run(text, owners, semitones=3)
        assert new_text == "Fm7"
        assert new_owners[-1] == 1

    def test_slash_bass_transposed(self):
        text = "G/B"
        owners = [0, 0, 1]
        new_text, _ = self._run(text, owners, semitones=2)
        assert new_text == "A/C#"

    def test_root_growth_absorbs_trailing_space(self):
        text = "C  G"
        owners = [0, 0, 0, 0]
        new_text, new_owners = self._run(text, owners, semitones=1, use_flats=True)
        assert new_text.startswith("Db")
        assert len(new_text) == len(new_owners)

    def test_root_shrink_pads_trailing_space(self):
        text = "C# G"
        owners = [0, 0, 0, 0]
        new_text, new_owners = self._run(text, owners, semitones=1, use_flats=False)
        assert len(new_text) == len(new_owners)

    def test_non_root_token_unchanged(self):
        text = "3 3"
        owners = [0, 0, 0]
        new_text, new_owners = self._run(text, owners)
        assert len(new_text) == len(new_owners)


# ---------------------------------------------------------------------------
# redistribute_to_runs
# ---------------------------------------------------------------------------


class TestRedistributeToRuns:
    def _make_run(self, text):
        run = MagicMock()
        run.text = text
        return run

    def _make_paragraph(self, texts):
        para = MagicMock()
        para.runs = [self._make_run(t_) for t_ in texts]
        return para

    def _redistribute(self, para, semitones=3, use_flats=False, german=False):
        spelling = t.choose_spelling(use_flats, german)
        t.redistribute_to_runs(para, semitones, spelling, german)

    def test_transposes_root_in_place(self):
        para = self._make_paragraph(["Am", "  ", "G"])
        self._redistribute(para, semitones=3)
        assert para.runs[0].text == "Cm"
        assert para.runs[2].text == "A#"

    def test_preserves_superscript_suffix_run(self):
        para = self._make_paragraph(["Dm", "7"])
        self._redistribute(para, semitones=3)
        assert para.runs[0].text == "Fm"
        assert para.runs[1].text == "7"

    def test_slash_bass_split_across_runs(self):
        para = self._make_paragraph(["G/", "B"])
        self._redistribute(para, semitones=2)
        assert para.runs[0].text == "A/"
        assert para.runs[1].text == "C#"

    def test_root_accidental_split_across_runs(self):
        para = self._make_paragraph(["A", "#"])
        self._redistribute(para, semitones=1, use_flats=False)
        assert (para.runs[0].text + para.runs[1].text).strip() == "B"

    def test_no_runs_does_nothing(self):
        para = MagicMock()
        para.runs = []
        self._redistribute(para)

    def test_empty_text_does_nothing(self):
        para = self._make_paragraph(["", ""])
        self._redistribute(para)
        assert para.runs[0].text == ""
        assert para.runs[1].text == ""


# ---------------------------------------------------------------------------
# _first_owner
# ---------------------------------------------------------------------------


class TestFirstOwner:
    def test_returns_first(self):
        assert t._first_owner([3, 1, 2]) == 3

    def test_defaults_to_zero_when_empty(self):
        assert t._first_owner([]) == 0


# ---------------------------------------------------------------------------
# superscript formatting preservation (integration)
# ---------------------------------------------------------------------------


class TestSuperscriptPreserved:
    def _make_superscript_docx(self):
        document = Document()
        paragraph = document.add_paragraph()
        base = paragraph.add_run("Dm")
        sup = paragraph.add_run("7")
        sup.font.superscript = True
        paragraph.add_run("         ")
        base2 = paragraph.add_run("G")
        sup2 = paragraph.add_run("sus-3")
        sup2.font.superscript = True
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue(), base, sup, base2, sup2

    def test_superscript_runs_survive_transposition(self):
        file_bytes, *_ = self._make_superscript_docx()

        docx_bytes, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "E")

        result = Document(BytesIO(docx_bytes))
        para = result.paragraphs[0]
        runs = [r for r in para.runs if r.text]

        assert runs[0].text == "F#m"
        assert runs[0].font.superscript is None
        assert runs[1].text == "7"
        assert runs[1].font.superscript is True
        assert runs[-2].text == "B"
        assert runs[-1].text == "sus-3"
        assert runs[-1].font.superscript is True
        assert ("Dm7", "F#m7") in changes
        assert ("Gsus-3", "Bsus-3") in changes


# ---------------------------------------------------------------------------
# process_paragraph
# ---------------------------------------------------------------------------


class TestProcessParagraph:
    def _make_paragraph(self, text):
        para = MagicMock()
        para.text = text
        run = MagicMock()
        run.text = text
        para.runs = [run]
        return para

    def test_chord_line_is_transposed(self):
        para = self._make_paragraph("C G Am F")
        changes = set()
        t.process_paragraph(para, 2, False, False, changes)
        assert len(changes) > 0

    def test_lyric_line_is_skipped(self):
        para = self._make_paragraph("Hello world this is lyrics")
        changes = set()
        t.process_paragraph(para, 2, False, False, changes)
        assert changes == set()

    def test_empty_paragraph_is_skipped(self):
        para = self._make_paragraph("")
        changes = set()
        t.process_paragraph(para, 2, False, False, changes)
        assert changes == set()


# ---------------------------------------------------------------------------
# iter_paragraphs
# ---------------------------------------------------------------------------


class TestIterParagraphs:
    def _make_document(self, body_texts, table_texts=None, header_texts=None, footer_texts=None):
        doc = MagicMock()

        doc.paragraphs = [self._make_para(t_) for t_ in body_texts]

        if table_texts:
            cell = MagicMock()
            cell.paragraphs = [self._make_para(t_) for t_ in table_texts]
            row = MagicMock()
            row.cells = [cell]
            table = MagicMock()
            table.rows = [row]
            doc.tables = [table]
        else:
            doc.tables = []

        section = MagicMock()
        section.header.paragraphs = [self._make_para(t_) for t_ in (header_texts or [])]
        section.footer.paragraphs = [self._make_para(t_) for t_ in (footer_texts or [])]
        doc.sections = [section]

        return doc

    def _make_para(self, text):
        para = MagicMock()
        para.text = text
        return para

    def test_yields_body_paragraphs(self):
        doc = self._make_document(["C G", "lyrics"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "C G" in texts
        assert "lyrics" in texts

    def test_yields_table_paragraphs(self):
        doc = self._make_document([], table_texts=["Am Dm"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "Am Dm" in texts

    def test_yields_header_paragraphs(self):
        doc = self._make_document([], header_texts=["F C"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "F C" in texts

    def test_yields_footer_paragraphs(self):
        doc = self._make_document([], footer_texts=["G D"])
        texts = [p.text for p in t.iter_paragraphs(doc)]
        assert "G D" in texts


# ---------------------------------------------------------------------------
# detect_german
# ---------------------------------------------------------------------------


class TestDetectGerman:
    def _make_document_with_para(self, text):
        para = MagicMock()
        para.text = text
        doc = MagicMock()
        doc.paragraphs = [para]
        doc.tables = []
        section = MagicMock()
        section.header.paragraphs = []
        section.footer.paragraphs = []
        doc.sections = [section]
        return doc

    def test_detects_german_notation(self):
        doc = self._make_document_with_para("H C D")
        assert t.detect_german(doc) is True

    def test_no_german_notation(self):
        doc = self._make_document_with_para("C G Am F")
        assert t.detect_german(doc) is False

    def test_lyric_line_with_h_word_ignored(self):
        doc = self._make_document_with_para("Hello world")
        assert t.detect_german(doc) is False


# ---------------------------------------------------------------------------
# find_input_file
# ---------------------------------------------------------------------------


class TestFindInputFile:
    def test_returns_none_when_no_docx(self, tmp_path):
        with patch.object(t, "INPUT_DIR", tmp_path):
            assert t.find_input_file() is None

    def test_returns_docx_file(self, tmp_path):
        docx_file = tmp_path / "song.docx"
        docx_file.touch()
        with patch.object(t, "INPUT_DIR", tmp_path):
            assert t.find_input_file() == docx_file

    def test_skips_temp_files(self, tmp_path):
        temp_file = tmp_path / "~$song.docx"
        temp_file.touch()
        with patch.object(t, "INPUT_DIR", tmp_path):
            assert t.find_input_file() is None


# ---------------------------------------------------------------------------
# _prompt_key
# ---------------------------------------------------------------------------


class TestPromptKey:
    def test_valid_input_returns_parsed_key(self):
        with patch("builtins.input", return_value="Am"):
            result = t._prompt_key("Key: ")
        assert result == ("A", True)

    def test_invalid_input_exits(self):
        with patch("builtins.input", return_value="X"), pytest.raises(SystemExit):
            t._prompt_key("Key: ")

    def test_invalid_input_prints_message(self, capsys):
        with patch("builtins.input", return_value="X"), pytest.raises(SystemExit):
            t._prompt_key("Key: ")
        assert "not a valid key" in capsys.readouterr().out


# ---------------------------------------------------------------------------
# _save_and_report
# ---------------------------------------------------------------------------


class TestSaveAndReport:
    def test_saves_file_and_prints_summary(self, tmp_path, capsys):
        doc = MagicMock()
        input_file = tmp_path / "song.docx"
        input_file.touch()
        output_dir = tmp_path / "output"

        with patch.object(t, "OUTPUT_DIR", output_dir):
            t._save_and_report(doc, input_file, "C", "D", {("C", "D"), ("G", "A")})

        captured = capsys.readouterr().out
        assert "C" in captured
        assert "D" in captured
        doc.save.assert_called_once()

    def test_reports_no_changes(self, tmp_path, capsys):
        doc = MagicMock()
        input_file = tmp_path / "song.docx"
        input_file.touch()
        output_dir = tmp_path / "output"

        with patch.object(t, "OUTPUT_DIR", output_dir):
            t._save_and_report(doc, input_file, "C", "C", set())

        assert "No chords were changed" in capsys.readouterr().out


# ---------------------------------------------------------------------------
# _transpose_document
# ---------------------------------------------------------------------------


class TestTransposeDocument:
    def test_returns_changes(self):
        para = MagicMock()
        para.text = "C G Am F"
        run = MagicMock()
        run.text = "C G Am F"
        para.runs = [run]

        doc = MagicMock()
        doc.paragraphs = [para]
        doc.tables = []
        section = MagicMock()
        section.header.paragraphs = []
        section.footer.paragraphs = []
        doc.sections = [section]

        changes = t._transpose_document(doc, 2, False, False)
        assert isinstance(changes, set)
        assert len(changes) > 0


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------


class TestMain:
    def _make_mock_document(self, para_text="C G Am F"):
        para = MagicMock()
        para.text = para_text
        run = MagicMock()
        run.text = para_text
        para.runs = [run]
        doc = MagicMock()
        doc.paragraphs = [para]
        doc.tables = []
        section = MagicMock()
        section.header.paragraphs = []
        section.footer.paragraphs = []
        doc.sections = [section]
        return doc

    def test_exits_when_no_input_file(self, tmp_path):
        with patch.object(t, "INPUT_DIR", tmp_path), pytest.raises(SystemExit):
            t.main()

    def test_full_run(self, tmp_path):
        input_dir = tmp_path / "input"
        input_dir.mkdir()
        output_dir = tmp_path / "output"
        docx_file = input_dir / "song.docx"
        docx_file.touch()

        mock_doc = self._make_mock_document()

        with (
            patch.object(t, "INPUT_DIR", input_dir),
            patch.object(t, "OUTPUT_DIR", output_dir),
            patch("transpose.Document", return_value=mock_doc),
            patch("builtins.input", side_effect=["C", "D"]),
        ):
            t.main()

        mock_doc.save.assert_called_once()


# ---------------------------------------------------------------------------
# transpose_document_bytes
# ---------------------------------------------------------------------------


def _make_docx_bytes(lines):
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestTransposeDocumentBytes:
    def test_transposes_and_returns_metadata(self):
        file_bytes = _make_docx_bytes(["C G Am F", "Hello world"])

        docx_bytes, from_label, to_label, changes = t.transpose_document_bytes(file_bytes, "C", "D")

        assert from_label == "C"
        assert to_label == "D"
        assert ("C", "D") in changes

        result = Document(BytesIO(docx_bytes))
        texts = [p.text for p in result.paragraphs]
        assert "Hello world" in texts

    def test_invalid_current_key_raises(self):
        file_bytes = _make_docx_bytes(["C G"])
        with pytest.raises(t.InvalidKeyError):
            t.transpose_document_bytes(file_bytes, "X", "D")

    def test_invalid_target_key_raises(self):
        file_bytes = _make_docx_bytes(["C G"])
        with pytest.raises(t.InvalidKeyError):
            t.transpose_document_bytes(file_bytes, "C", "X")

    def test_same_key_produces_no_changes(self):
        file_bytes = _make_docx_bytes(["C G Am F"])
        _, from_label, to_label, changes = t.transpose_document_bytes(file_bytes, "C", "C")
        assert from_label == "C"
        assert to_label == "C"
        assert changes == []

    def test_minor_key_labels(self):
        file_bytes = _make_docx_bytes(["Am Dm E"])
        _, from_label, to_label, _ = t.transpose_document_bytes(file_bytes, "Am", "Bm")
        assert from_label == "Am"
        assert to_label == "Bm"

    def test_german_target_key_on_english_document(self):
        file_bytes = _make_docx_bytes(["C G Am F"])
        _, from_label, to_label, changes = t.transpose_document_bytes(file_bytes, "C", "Es")
        assert from_label == "C"
        assert to_label == "Es"
        assert ("C", "Eb") in changes

    def test_german_current_key_on_english_document(self):
        file_bytes = _make_docx_bytes(["C G Am F"])
        _, from_label, to_label, _ = t.transpose_document_bytes(file_bytes, "Cis", "C")
        assert from_label == "Cis"
        assert to_label == "C"


# ---------------------------------------------------------------------------
# German note handling
# ---------------------------------------------------------------------------


class TestGermanNotes:
    def test_es_is_e_flat(self):
        assert t.note_semitone("Es", True) == 3

    def test_as_is_a_flat(self):
        assert t.note_semitone("As", True) == 8

    def test_german_b_is_b_flat(self):
        assert t.note_semitone("B", True) == 10

    def test_cis_is_c_sharp(self):
        assert t.note_semitone("Cis", True) == 1

    def test_transpose_es_up_two_uses_german_spelling(self):
        assert t.transpose_note("Es", 2, t.GERMAN_FLAT_SPELLING, True) == "F"

    def test_transpose_into_flat_produces_german_name(self):
        assert t.transpose_note("C", 3, t.GERMAN_FLAT_SPELLING, True) == "Es"

    def test_transpose_into_sharp_produces_german_name(self):
        assert t.transpose_note("C", 1, t.GERMAN_SHARP_SPELLING, True) == "Cis"

    def test_transpose_chord_german_slash(self):
        result = t.transpose_chord("B/d", 2, t.GERMAN_FLAT_SPELLING, True)
        assert result == "C/E"

    def test_line_uses_german_detects_es(self):
        assert t.line_uses_german("Es gm B F") is True

    def test_line_uses_german_detects_as(self):
        assert t.line_uses_german("As Des") is True

    def test_line_without_german_names(self):
        assert t.line_uses_german("C G Am F") is False


# ---------------------------------------------------------------------------
# split_affixes / punctuation handling
# ---------------------------------------------------------------------------


class TestSplitAffixes:
    def test_no_affixes(self):
        assert t.split_affixes("Am") == ("", "Am", "")

    def test_trailing_ellipsis(self):
        assert t.split_affixes("B…") == ("", "B", "…")

    def test_wrapping_parentheses(self):
        assert t.split_affixes("(C)") == ("(", "C", ")")

    def test_all_punctuation(self):
        assert t.split_affixes("...") == ("...", "", "")


class TestPunctuationTransposition:
    def test_chord_with_trailing_ellipsis(self):
        assert t.transpose_chord("B…", 1, t.SHARP_SPELLING, False) == "C…"

    def test_non_chord_core_left_untouched(self):
        assert t.transpose_chord("(the)", 2, t.SHARP_SPELLING, False) == "(the)"


# ---------------------------------------------------------------------------
# is_chord_token / is_chord_line decoration handling
# ---------------------------------------------------------------------------


class TestChordLineDecorations:
    def test_chord_with_punctuation_is_chord_token(self):
        assert t.is_chord_token("B…") is True

    def test_word_is_not_chord_token(self):
        assert t.is_chord_token("the") is False

    def test_line_with_label_and_chords(self):
        assert t.is_chord_line("(Vége: Es B/d F/c B )") is True

    def test_line_with_section_label(self):
        assert t.is_chord_line("REF: C G Am") is True

    def test_lyric_line_with_chord_word_rejected(self):
        assert t.is_chord_line("C the G") is False

    def test_pure_lyric_line_rejected(self):
        assert t.is_chord_line("Nem számít hol jártam") is False

    def test_line_with_no_chords_rejected(self):
        assert t.is_chord_line("(Vége: )") is False


class TestChordLineWithOwnersPunctuation:
    def _run(self, text, semitones=1, use_flats=False, german=False):
        owners = [0] * len(text)
        spelling = t.choose_spelling(use_flats, german)
        return t.transpose_line_with_owners(text, owners, semitones, spelling, german)

    def test_transposes_chords_keeps_decorations(self):
        new_text, new_owners = self._run("(Vége: B )", german=True)
        assert new_text == "(Vége: H )"
        assert len(new_text) == len(new_owners)

    def test_trailing_ellipsis_preserved(self):
        new_text, new_owners = self._run("B…")
        assert new_text == "C…"
        assert len(new_text) == len(new_owners)

    def test_non_chord_word_left_untouched(self):
        new_text, _ = self._run("the")
        assert new_text == "the"


class TestTransposeTokenWithOwnersDirect:
    def test_non_root_token_returns_unchanged(self):
        owners = [0, 0]
        result, result_owners = t._transpose_token_with_owners(
            "3x", owners, 2, t.SHARP_SPELLING, False
        )
        assert result == "3x"
        assert result_owners == owners

    def test_punctuated_non_chord_returns_unchanged(self):
        owners = [0, 0, 0, 0, 0]
        result, result_owners = t._transpose_token_with_owners(
            "(the)", owners, 2, t.SHARP_SPELLING, False
        )
        assert result == "(the)"
        assert result_owners == owners


class TestTransposeLineTextNonChord:
    def test_non_chord_tokens_are_preserved(self):
        spelling = t.choose_spelling(False, False)
        changes = set()
        result = t.transpose_line_text("(Vége: C )", 2, spelling, False, changes)
        assert result == "(Vége: D )"
        assert ("C", "D") in changes


# ---------------------------------------------------------------------------
# convert_docx_to_pdf
# ---------------------------------------------------------------------------


class TestConvertDocxToPdf:
    def test_returns_pdf_bytes_on_success(self):
        response = MagicMock(status_code=200, content=b"%PDF-1.7")
        with patch("transpose.requests.post", return_value=response) as post:
            result = t.convert_docx_to_pdf(b"docx")
        assert result == b"%PDF-1.7"
        endpoint = post.call_args.args[0]
        assert endpoint.endswith("/forms/libreoffice/convert")

    def test_uses_configured_gotenberg_url(self):
        response = MagicMock(status_code=200, content=b"%PDF-1.7")
        with (
            patch.dict("os.environ", {"GOTENBERG_URL": "http://gotenberg:3000/"}),
            patch("transpose.requests.post", return_value=response) as post,
        ):
            t.convert_docx_to_pdf(b"docx")
        assert post.call_args.args[0] == "http://gotenberg:3000/forms/libreoffice/convert"

    def test_raises_on_timeout(self):
        with patch("transpose.requests.post", side_effect=requests.Timeout()):
            with pytest.raises(t.PdfConversionError, match="timed out"):
                t.convert_docx_to_pdf(b"docx")

    def test_raises_on_connection_error(self):
        with patch("transpose.requests.post", side_effect=requests.ConnectionError()):
            with pytest.raises(t.PdfConversionError, match="reach"):
                t.convert_docx_to_pdf(b"docx")

    def test_raises_on_non_200_response(self):
        response = MagicMock(status_code=500, content=b"error")
        with patch("transpose.requests.post", return_value=response):
            with pytest.raises(t.PdfConversionError, match="failed"):
                t.convert_docx_to_pdf(b"docx")


class TestGotenbergUrl:
    def test_defaults_to_localhost(self):
        with patch.dict("os.environ", {}, clear=True):
            assert t._gotenberg_url() == "http://localhost:3000"

    def test_reads_env_and_strips_trailing_slash(self):
        with patch.dict("os.environ", {"GOTENBERG_URL": "http://gotenberg:3000/"}):
            assert t._gotenberg_url() == "http://gotenberg:3000"
