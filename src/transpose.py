"""Transpose chord sheets in .docx format between musical keys."""

import os
import re
import sys
from io import BytesIO
from pathlib import Path

import requests
from docx import Document

DEFAULT_GOTENBERG_URL = "http://localhost:3000"
PDF_CONVERSION_TIMEOUT_SECONDS = 120

PROJECT_ROOT = Path(__file__).parent.parent
INPUT_DIR = PROJECT_ROOT / "input"
OUTPUT_DIR = PROJECT_ROOT / "output"

NOTE_TO_SEMITONE = {
    "C": 0,
    "C#": 1,
    "Db": 1,
    "D": 2,
    "D#": 3,
    "Eb": 3,
    "E": 4,
    "Fb": 4,
    "E#": 5,
    "F": 5,
    "F#": 6,
    "Gb": 6,
    "G": 7,
    "G#": 8,
    "Ab": 8,
    "A": 9,
    "A#": 10,
    "Bb": 10,
    "B": 11,
    "Cb": 11,
    "B#": 0,
    "H": 11,
    "H#": 0,
}

GERMAN_NOTE_NAMES = {
    "C": 0,
    "Cis": 1,
    "Des": 1,
    "D": 2,
    "Dis": 3,
    "Es": 3,
    "Eis": 5,
    "E": 4,
    "Fes": 4,
    "F": 5,
    "Fis": 6,
    "Ges": 6,
    "G": 7,
    "Gis": 8,
    "As": 8,
    "A": 9,
    "Ais": 10,
    "B": 10,
    "His": 0,
    "H": 11,
    "Ces": 11,
}

GERMAN_NOTE_TO_SEMITONE = dict(NOTE_TO_SEMITONE)
GERMAN_NOTE_TO_SEMITONE["B"] = 10
GERMAN_NOTE_TO_SEMITONE["Bb"] = 9
GERMAN_NOTE_TO_SEMITONE.update(GERMAN_NOTE_NAMES)

GERMAN_ROOT_NAMES = sorted(GERMAN_NOTE_NAMES, key=len, reverse=True)
GERMAN_ONLY_ROOTS_KEYS = {name for name in GERMAN_NOTE_NAMES if name not in NOTE_TO_SEMITONE}

SHARP_KEYS = {"G", "D", "A", "E", "B", "F#", "C#", "Em", "Bm", "F#m", "C#m", "G#m", "D#m", "A#m"}
FLAT_KEYS = {"F", "Bb", "Eb", "Ab", "Db", "Gb", "Cb", "Dm", "Gm", "Cm", "Fm", "Bbm", "Ebm", "Abm"}

SHARP_SPELLING = ["C", "C#", "D", "D#", "E", "F", "F#", "G", "G#", "A", "A#", "B"]
FLAT_SPELLING = ["C", "Db", "D", "Eb", "E", "F", "Gb", "G", "Ab", "A", "Bb", "B"]

GERMAN_SHARP_SPELLING = ["C", "Cis", "D", "Dis", "E", "F", "Fis", "G", "Gis", "A", "B", "H"]
GERMAN_FLAT_SPELLING = ["C", "Des", "D", "Es", "E", "F", "Ges", "G", "As", "A", "B", "H"]

ROOT_PATTERN = r"(?:Cis|Des|Dis|Eis|Fis|Ges|Gis|Ais|His|Ces|Fes|Es|As|[A-Ha-h](?:#|b)?)"
QUALITY_PATTERN = r"(?:maj|min|dim|aug|sus|add|m|M|\+|°|ø|\d|#|b|-|\(|\)|,)*"
CHORD_TOKEN_RE = re.compile(
    rf"^{ROOT_PATTERN}{QUALITY_PATTERN}(?:/{ROOT_PATTERN}{QUALITY_PATTERN})?$"
)
ROOT_MATCH_RE = re.compile(rf"^({ROOT_PATTERN})(.*)$")
TOKEN_SPLIT_RE = re.compile(r"(\S+)(\s*)")


def _is_minor_suffix(rest):
    """Return True if the trailing text denotes a minor key."""
    return rest.strip().lower() in ("m", "min", "minor")


def parse_key(raw):
    """Parse a raw key string into a (note, is_minor) tuple, or None if invalid."""
    text = raw.strip()
    if not text:
        return None

    for name in GERMAN_ROOT_NAMES:
        if len(name) > 1 and text[: len(name)].capitalize() == name:
            rest = text[len(name) :]
            if rest == "" or _is_minor_suffix(rest):
                return name, _is_minor_suffix(rest)

    root = text[0].upper()
    rest = text[1:]
    accidental = ""
    if rest and rest[0] in "#b":
        accidental = rest[0]
        rest = rest[1:]
    note = root + accidental
    if note not in NOTE_TO_SEMITONE:
        return None
    return note, _is_minor_suffix(rest)


def key_label(note, minor):
    """Return the key label string, appending 'm' for minor keys."""
    return note + ("m" if minor else "")


def prefers_flats(target_note, target_minor):
    """Return True if the target key conventionally uses flat spelling."""
    label = key_label(target_note, target_minor)
    if label in FLAT_KEYS or target_note in FLAT_KEYS:
        return True
    if label in SHARP_KEYS or target_note in SHARP_KEYS:
        return False
    return True


def choose_spelling(use_flats, german):
    """Return the appropriate chromatic spelling list for the given preferences."""
    if german:
        return GERMAN_FLAT_SPELLING if use_flats else GERMAN_SHARP_SPELLING
    return FLAT_SPELLING if use_flats else SHARP_SPELLING


def note_semitone(note, german):
    """Return the semitone index (0–11) for a note name."""
    key = note[0].upper() + note[1:]
    table = GERMAN_NOTE_TO_SEMITONE if german else NOTE_TO_SEMITONE
    return table[key]


def key_semitone(note, german):
    """Return the semitone index for a parsed key note.

    German-only names such as Es or Cis resolve via the German table regardless
    of the document flag, while ambiguous names (e.g. B) keep their meaning for
    the active notation so an English key B stays B natural.
    """
    canonical = note[0].upper() + note[1:]
    if canonical in GERMAN_ONLY_ROOTS_KEYS:
        return GERMAN_NOTE_TO_SEMITONE[canonical]
    return note_semitone(note, german)


def transpose_note(note, semitones, spelling, german):
    """Transpose a single note name by the given number of semitones."""
    semitone = (note_semitone(note, german) + semitones) % 12
    return spelling[semitone]


PUNCTUATION = "().,;:…\"'-"


def split_affixes(token):
    """Split a token into (leading_punctuation, core, trailing_punctuation)."""
    start = 0
    while start < len(token) and token[start] in PUNCTUATION:
        start += 1
    end = len(token)
    while end > start and token[end - 1] in PUNCTUATION:
        end -= 1
    return token[:start], token[start:end], token[end:]


def transpose_chord(chord, semitones, spelling, german):
    """Transpose a chord symbol (including optional bass note) by semitones."""
    leading, core, trailing = split_affixes(chord)
    if core != chord:
        if not CHORD_TOKEN_RE.match(core):
            return chord
        return leading + transpose_chord(core, semitones, spelling, german) + trailing

    match = ROOT_MATCH_RE.match(chord)
    if not match:
        return chord
    root, remainder = match.group(1), match.group(2)
    new_root = transpose_note(root, semitones, spelling, german)
    if "/" in remainder:
        before_slash, bass = remainder.split("/", 1)
        bass_match = ROOT_MATCH_RE.match(bass)
        if bass_match:
            new_bass = transpose_note(bass_match.group(1), semitones, spelling, german)
            return new_root + before_slash + "/" + new_bass + bass_match.group(2)
    return new_root + remainder


def is_chord_token(token):
    """Return True if the bare token (ignoring surrounding punctuation) is a chord symbol."""
    core = token.strip("().,;:…\"'-")
    return bool(core) and bool(CHORD_TOKEN_RE.match(core))


def _is_decoration(token):
    """Return True if a non-chord token is punctuation or a section label, not a lyric word."""
    _, core, _ = split_affixes(token)
    if core == "":
        return True
    if core.endswith(":") or ":" in token:
        return True
    return core.isupper()


def is_chord_line(text):
    """Return True if the line is a chord line rather than lyrics.

    A line qualifies when every token is a chord symbol, or when it contains at
    least one chord and every remaining token is a decoration (punctuation or a
    section label). This tolerates lines like '(Vége: Es B/d F/c B )' while
    still rejecting lyric lines that merely contain a stray chord-like word.
    """
    tokens = text.split()
    if not tokens:
        return False
    chord_count = 0
    for token in tokens:
        if is_chord_token(token):
            chord_count += 1
        elif not _is_decoration(token):
            return False
    return chord_count > 0


GERMAN_ONLY_ROOTS = {name.upper() for name in GERMAN_NOTE_NAMES if len(name) > 1}
GERMAN_ONLY_ROOTS.add("H")


def _token_roots(token):
    """Yield each note root (main and bass) found in a chord token."""
    _, core, _ = split_affixes(token)
    for part in core.split("/"):
        match = ROOT_MATCH_RE.match(part)
        if match:
            yield match.group(1)


def line_uses_german(text):
    """Return True if the line contains an unambiguously German note name (H, Es, As, Cis…)."""
    for token in text.split():
        for root in _token_roots(token):
            if root.upper() in GERMAN_ONLY_ROOTS:
                return True
    return False


def transpose_line_text(text, semitones, spelling, german, changes):
    """Transpose all chord tokens in a line, preserving spacing, and record changes."""
    leading = text[: len(text) - len(text.lstrip())]
    body = text[len(leading) :]
    result = leading

    for token_match in TOKEN_SPLIT_RE.finditer(body):
        chord = token_match.group(1)
        trailing = token_match.group(2)
        if is_chord_token(chord):
            transposed = transpose_chord(chord, semitones, spelling, german)
        else:
            transposed = chord
        if transposed != chord:
            changes.add((chord, transposed))

        diff = len(transposed) - len(chord)
        if diff > 0:
            trailing = trailing[min(diff, len(trailing)) :]
        elif diff < 0:
            trailing = trailing + " " * (-diff)
        result += transposed + trailing

    return result


def _build_owner_map(runs):
    """Return the concatenated run text and a list mapping each char to its run index."""
    text = ""
    owners = []
    for run_index, run in enumerate(runs):
        text += run.text
        owners.extend(run_index for _ in run.text)
    return text, owners


def _transpose_token_with_owners(chord, owners, semitones, spelling, german):
    """Transpose a chord and return (new_chord, new_owners) preserving per-char runs.

    Owners for the root and bass note come from the original note's first character,
    so the surrounding formatting (e.g. superscript suffix runs) is kept intact.
    """
    leading, core, trailing = split_affixes(chord)
    if core != chord:
        if not CHORD_TOKEN_RE.match(core):
            return chord, owners
        core_owners = owners[len(leading) : len(leading) + len(core)]
        new_core, new_core_owners = _transpose_token_with_owners(
            core, core_owners, semitones, spelling, german
        )
        new_chord = leading + new_core + trailing
        new_owners = owners[: len(leading)] + new_core_owners + owners[len(leading) + len(core) :]
        return new_chord, new_owners

    match = ROOT_MATCH_RE.match(chord)
    if not match:
        return chord, owners

    root, remainder = match.group(1), match.group(2)
    root_owner = owners[0]
    new_root = transpose_note(root, semitones, spelling, german)
    new_chord = new_root
    new_owners = [root_owner] * len(new_root)
    remainder_owners = owners[len(root) :]

    if "/" in remainder:
        before_slash, bass = remainder.split("/", 1)
        bass_match = ROOT_MATCH_RE.match(bass)
        if bass_match:
            bass_root = bass_match.group(1)
            split_at = len(before_slash) + 1
            slash_part = remainder[:split_at]
            new_chord += slash_part
            new_owners += remainder_owners[:split_at]

            bass_owner = remainder_owners[split_at]
            new_bass = transpose_note(bass_root, semitones, spelling, german)
            new_chord += new_bass
            new_owners += [bass_owner] * len(new_bass)

            tail = remainder[split_at + len(bass_root) :]
            new_chord += tail
            new_owners += remainder_owners[split_at + len(bass_root) :]
            return new_chord, new_owners

    new_chord += remainder
    new_owners += remainder_owners
    return new_chord, new_owners


def transpose_line_with_owners(text, owners, semitones, spelling, german, changes=None):
    """Transpose a chord line, returning (new_text, new_owners) aligned to source runs.

    When a ``changes`` set is supplied, every ``(original, transposed)`` chord
    pair is recorded so callers get the change summary from this single path.
    """
    leading_len = len(text) - len(text.lstrip())
    result = text[:leading_len]
    result_owners = list(owners[:leading_len])

    body = text[leading_len:]
    body_owners = owners[leading_len:]

    cursor = 0
    for token_match in TOKEN_SPLIT_RE.finditer(body):
        chord = token_match.group(1)
        trailing = token_match.group(2)
        start = token_match.start(1)
        chord_owners = body_owners[start : start + len(chord)]
        trailing_owners = body_owners[start + len(chord) : token_match.end(2)]

        if is_chord_token(chord):
            new_chord, new_chord_owners = _transpose_token_with_owners(
                chord, chord_owners, semitones, spelling, german
            )
        else:
            new_chord, new_chord_owners = chord, list(chord_owners)

        if changes is not None and new_chord != chord:
            changes.add((chord, new_chord))

        diff = len(new_chord) - len(chord)
        if diff > 0:
            drop = min(diff, len(trailing))
            trailing = trailing[drop:]
            trailing_owners = trailing_owners[drop:]
        elif diff < 0:
            pad_owner = new_chord_owners[-1] if new_chord_owners else _first_owner(body_owners)
            trailing += " " * (-diff)
            trailing_owners = list(trailing_owners) + [pad_owner] * (-diff)

        result += new_chord + trailing
        result_owners += new_chord_owners + list(trailing_owners)
        cursor = token_match.end(2)

    result += body[cursor:]
    result_owners += list(body_owners[cursor:])
    return result, result_owners


def _first_owner(owners):
    """Return the first owner index, defaulting to 0 for empty owner lists."""
    return owners[0] if owners else 0


def redistribute_to_runs(paragraph, semitones, spelling, german, changes=None):
    """Transpose chord roots and write text back into runs, preserving formatting."""
    runs = paragraph.runs
    if not runs:
        return

    text, owners = _build_owner_map(runs)
    if not text:
        return

    new_text, new_owners = transpose_line_with_owners(
        text, owners, semitones, spelling, german, changes
    )

    run_texts = ["" for _ in runs]
    for char, owner in zip(new_text, new_owners, strict=True):
        run_texts[owner] += char

    for run, run_text in zip(runs, run_texts, strict=True):
        if run.text != run_text:
            run.text = run_text


def process_paragraph(paragraph, semitones, use_flats, german, changes):
    """Transpose all chords in a paragraph if it is a chord line."""
    text = paragraph.text
    if not is_chord_line(text):
        return
    spelling = choose_spelling(use_flats, german)
    redistribute_to_runs(paragraph, semitones, spelling, german, changes)


def iter_paragraphs(document):
    """Yield every paragraph in the document, including tables, headers, and footers."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs


def detect_german(document):
    """Return True if any chord line in the document uses German note notation."""
    for paragraph in iter_paragraphs(document):
        text = paragraph.text
        if is_chord_line(text) and line_uses_german(text):
            return True
    return False


def find_input_file():
    """Return the first .docx file in INPUT_DIR, or None if none exists."""
    docx_files = [p for p in INPUT_DIR.glob("*.docx") if not p.name.startswith("~$")]
    if not docx_files:
        return None
    return docx_files[0]


def _prompt_key(prompt_text):
    """Prompt the user for a key string, parse it, and exit on invalid input."""
    raw = input(prompt_text)
    parsed = parse_key(raw)
    if parsed is None:
        print(f"'{raw.strip()}' is not a valid key.")
        sys.exit(1)
    return parsed


def _transpose_document(document, semitones, use_flats, german):
    """Apply transposition to every paragraph in the document and return the change set."""
    changes: set[tuple[str, str]] = set()
    for paragraph in iter_paragraphs(document):
        process_paragraph(paragraph, semitones, use_flats, german, changes)
    return changes


def _save_and_report(document, input_file, from_label, to_label, changes):
    """Save the transposed document and print a summary of chord changes."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    output_file = OUTPUT_DIR / f"{input_file.stem}_{to_label}.docx"
    document.save(output_file)

    print(f"\nSheet transposed from {from_label} to {to_label}\n")
    if changes:
        for original, transposed in sorted(changes):
            print(f"{original} -> {transposed}")
    else:
        print("No chords were changed.")
    print(f"\nSaved to: {output_file}")


class InvalidKeyError(ValueError):
    """Raised when a supplied key string cannot be parsed."""


class PdfConversionError(RuntimeError):
    """Raised when a .docx cannot be converted to PDF."""


def _gotenberg_url():
    """Return the base URL of the Gotenberg conversion service."""
    return (os.environ.get("GOTENBERG_URL") or DEFAULT_GOTENBERG_URL).rstrip("/")


def convert_docx_to_pdf(docx_bytes):
    """Convert .docx bytes to PDF bytes via the Gotenberg conversion service."""
    endpoint = f"{_gotenberg_url()}/forms/libreoffice/convert"
    files = {
        "files": (
            "sheet.docx",
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    try:
        response = requests.post(endpoint, files=files, timeout=PDF_CONVERSION_TIMEOUT_SECONDS)
    except requests.Timeout as exc:
        raise PdfConversionError("PDF conversion timed out.") from exc
    except requests.RequestException as exc:
        raise PdfConversionError("Could not reach the PDF conversion service.") from exc

    if response.status_code != 200:
        raise PdfConversionError("The PDF conversion service failed to convert the document.")

    return response.content


def _text_uses_german(text):
    """Return True if any chord line in the text uses German note notation."""
    for line in text.splitlines():
        if is_chord_line(line) and line_uses_german(line):
            return True
    return False


def transpose_text(text, current_key, target_key):
    """Transpose chord lines in plain text and return (text, from_label, to_label, changes).

    Only lines detected as chord lines are transposed; lyric and blank lines are
    left untouched. Notation, spelling, and German H handling match the .docx path.
    """
    current = parse_key(current_key)
    if current is None:
        raise InvalidKeyError(f"'{current_key}' is not a valid key.")
    target = parse_key(target_key)
    if target is None:
        raise InvalidKeyError(f"'{target_key}' is not a valid key.")

    current_note, current_minor = current
    target_note, target_minor = target

    german = _text_uses_german(text)
    semitones = (key_semitone(target_note, german) - key_semitone(current_note, german)) % 12
    use_flats = prefers_flats(target_note, target_minor)
    spelling = choose_spelling(use_flats, german)

    changes: set[tuple[str, str]] = set()
    result_parts = []
    for segment in text.splitlines(keepends=True):
        content = segment.splitlines()[0]
        ending = segment[len(content) :]
        if is_chord_line(content):
            content = transpose_line_text(content, semitones, spelling, german, changes)
        result_parts.append(content + ending)

    result = "".join(result_parts)

    from_label = key_label(current_note, current_minor)
    to_label = key_label(target_note, target_minor)

    return result, from_label, to_label, sorted(changes)


def transpose_document_bytes(file_bytes, current_key, target_key):
    """Transpose a .docx given as bytes and return (docx_bytes, from_label, to_label, changes)."""
    current = parse_key(current_key)
    if current is None:
        raise InvalidKeyError(f"'{current_key}' is not a valid key.")
    target = parse_key(target_key)
    if target is None:
        raise InvalidKeyError(f"'{target_key}' is not a valid key.")

    current_note, current_minor = current
    target_note, target_minor = target

    document = Document(BytesIO(file_bytes))
    german = detect_german(document)

    semitones = (key_semitone(target_note, german) - key_semitone(current_note, german)) % 12
    use_flats = prefers_flats(target_note, target_minor)

    changes = _transpose_document(document, semitones, use_flats, german)

    from_label = key_label(current_note, current_minor)
    to_label = key_label(target_note, target_minor)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue(), from_label, to_label, sorted(changes)


def main():
    """Entry point: load the input file, prompt for keys, transpose, and save."""
    input_file = find_input_file()
    if input_file is None:
        print(f"No .docx file found in {INPUT_DIR}. Drop your sheet there first.")
        sys.exit(1)

    print(f"Found input file: {input_file.name}\n")

    current_note, current_minor = _prompt_key("Current key of the sheet (e.g. C, Am, F#): ")
    target_note, target_minor = _prompt_key("Desired key to transpose to (e.g. D, Bm, Eb): ")

    document = Document(str(input_file))
    german = detect_german(document)

    semitones = (key_semitone(target_note, german) - key_semitone(current_note, german)) % 12
    use_flats = prefers_flats(target_note, target_minor)

    changes = _transpose_document(document, semitones, use_flats, german)

    from_label = key_label(current_note, current_minor)
    to_label = key_label(target_note, target_minor)

    _save_and_report(document, input_file, from_label, to_label, changes)


if __name__ == "__main__":
    main()
