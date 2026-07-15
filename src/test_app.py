"""Tests for the Flask web layer in app.py."""
# pylint: disable=redefined-outer-name,missing-function-docstring

import json
import re
from io import BytesIO
from urllib.parse import unquote

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter

import app as app_module
import seo
from app import app

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _one_page_pdf():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture
def client():
    app.config["TESTING"] = True
    with app.test_client() as test_client:
        yield test_client


def _chord_docx_bytes():
    document = Document()
    document.add_paragraph("C G Am F")
    document.add_paragraph("These are the lyrics")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _upload(data_bytes, filename="song.docx") -> dict[str, object]:
    return {"file": (BytesIO(data_bytes), filename)}


def test_index_lists_keys(client):
    response = client.get("/")
    assert response.status_code == 200
    assert b"Chord Transposer" in response.data


def test_index_shows_docx_formatting_moat_badge(client):
    body = client.get("/").get_data(as_text=True)
    assert "Keeps your Word (.docx) formatting intact" in body


def test_index_links_versioned_stylesheet(client):
    response = client.get("/")
    body = response.get_data(as_text=True)
    assert "cdn.tailwindcss.com" not in body
    assert f"/static/tailwind.css?v={app_module.CSS_VERSION}" in body


def test_index_html_is_not_cached(client):
    response = client.get("/")
    assert response.headers["Cache-Control"] == "no-cache"


def test_static_css_is_cached_immutably(client):
    response = client.get(f"/static/tailwind.css?v={app_module.CSS_VERSION}")
    assert response.status_code == 200
    cache_control = response.headers["Cache-Control"]
    assert "immutable" in cache_control
    assert f"max-age={app_module.STATIC_CSS_MAX_AGE}" in cache_control


def test_css_version_is_stable_hash():
    assert app_module.CSS_VERSION == app_module._compute_css_version()
    assert len(app_module.CSS_VERSION) == 12


def test_css_version_falls_back_when_missing(monkeypatch):
    def _raise(_self):
        raise OSError("missing")

    monkeypatch.setattr(app_module.Path, "read_bytes", _raise)
    assert app_module._compute_css_version() == "dev"


def test_index_references_hashed_js_bundle(client):
    response = client.get("/")
    body = response.get_data(as_text=True)
    assert f'src="{app_module.JS_BUNDLE}"' in body
    assert app_module.JS_BUNDLE.startswith("/static/js/")


def test_js_bundle_reads_hashed_entry_from_manifest(monkeypatch):
    monkeypatch.setattr(
        app_module.Path,
        "read_text",
        lambda _self: '{"main.js": {"file": "main.abc123.js"}}',
    )
    assert app_module._resolve_js_bundle() == "/static/js/main.abc123.js"


def test_js_bundle_falls_back_when_manifest_missing(monkeypatch):
    def _raise(_self):
        raise OSError("missing")

    monkeypatch.setattr(app_module.Path, "read_text", _raise)
    assert app_module._resolve_js_bundle() == app_module.JS_BUNDLE_FALLBACK


def test_js_bundle_falls_back_when_entry_malformed(monkeypatch):
    monkeypatch.setattr(app_module.Path, "read_text", lambda _self: "{}")
    assert app_module._resolve_js_bundle() == app_module.JS_BUNDLE_FALLBACK


def test_health(client):
    response = client.get("/health")
    assert response.status_code == 200
    assert response.get_json() == {"status": "ok"}


def test_robots_txt(client):
    response = client.get("/robots.txt")
    assert response.status_code == 200
    assert response.mimetype == "text/plain"
    assert "Sitemap:" in response.get_data(as_text=True)


def test_og_image_png(client):
    response = client.get("/og-image.png")
    assert response.status_code == 200
    assert response.mimetype == "image/png"
    assert response.get_data().startswith(b"\x89PNG")


def test_favicon_svg(client):
    response = client.get("/favicon.svg")
    assert response.status_code == 200
    assert response.mimetype == "image/svg+xml"
    assert response.get_data(as_text=True).lstrip().startswith("<svg")


def test_sitemap_xml(client):
    response = client.get("/sitemap.xml")
    assert response.status_code == 200
    assert response.mimetype == "application/xml"
    assert "<urlset" in response.get_data(as_text=True)


def test_transpose_missing_file(client):
    response = client.post(
        "/transpose",
        data={"current_key": "C", "target_key": "D"},
        content_type="multipart/form-data",
    )
    assert response.status_code == 400
    assert "choose a .docx" in response.get_json()["error"]


def test_transpose_wrong_extension(client):
    data = _upload(b"not a docx", filename="song.txt")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Only .docx" in response.get_json()["error"]


def test_transpose_missing_keys(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "", "target_key": ""})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "keys are required" in response.get_json()["error"]


def test_transpose_unsupported_format(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D", "format": "txt"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Unsupported output format" in response.get_json()["error"]


def test_transpose_empty_file(client):
    data = _upload(b"", filename="empty.docx")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "empty" in response.get_json()["error"]


def test_transpose_invalid_key(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "X", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "not a valid key" in response.get_json()["error"]


def test_transpose_invalid_docx(client):
    data = _upload(b"this is not a real docx", filename="bad.docx")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "valid .docx" in response.get_json()["error"]


def test_transpose_docx_success(client):
    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 200
    assert response.mimetype == DOCX_MIMETYPE
    assert "song_D.docx" in response.headers["Content-Disposition"]
    assert unquote(response.headers["X-Transpose-From"]) == "C"
    assert unquote(response.headers["X-Transpose-To"]) == "D"
    assert "C->D" in unquote(response.headers["X-Transpose-Changes"])


def test_transpose_preserves_non_ascii_filename(client):
    data = _upload(_chord_docx_bytes(), filename="Covered-Gizus_rövid_G.docx")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 200
    disposition = response.headers["Content-Disposition"]
    assert "filename*=UTF-8''Covered-Gizus_r%C3%B6vid_G_D.docx" in disposition


def test_transpose_pdf_success(client, monkeypatch):
    monkeypatch.setattr(app_module, "convert_docx_to_pdf", lambda docx_bytes: b"%PDF-1.4 fake")

    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D", "format": "pdf"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 200
    assert response.mimetype == "application/pdf"
    assert "song_D.pdf" in response.headers["Content-Disposition"]
    assert response.data == b"%PDF-1.4 fake"


def test_transpose_pdf_conversion_failure(client, monkeypatch):
    def _raise(_docx_bytes):
        raise app_module.PdfConversionError("PDF conversion timed out.")

    monkeypatch.setattr(app_module, "convert_docx_to_pdf", _raise)

    data = _upload(_chord_docx_bytes())
    data.update({"current_key": "C", "target_key": "D", "format": "pdf"})
    response = client.post("/transpose", data=data, content_type="multipart/form-data")

    assert response.status_code == 503
    assert "timed out" in response.get_json()["error"]


def _jsonld_blocks(body):
    pattern = r'type="application/ld\+json">(.*?)</script>'
    return [json.loads(block) for block in re.findall(pattern, body, re.S)]


def _jsonld_types(body):
    return {block.get("@type") for block in _jsonld_blocks(body)}


def _title(body):
    return re.search(r"<title>(.*?)</title>", body, re.S).group(1).strip()


def _canonical(body):
    return re.search(r'<link rel="canonical" href="(.*?)"', body).group(1)


def test_landing_pages_return_200_with_unique_title_and_canonical(client):
    titles = {}
    canonicals = {}
    for page in seo.landing_pages():
        response = client.get(page["path"])
        assert response.status_code == 200, page["path"]
        body = response.get_data(as_text=True)
        titles.setdefault(_title(body), []).append(page["path"])
        canonicals.setdefault(_canonical(body), []).append(page["path"])
        assert _canonical(body) == f"{seo.SITE_URL}{page['path']}"

    landing_count = len(seo.landing_pages())
    assert len(titles) == landing_count
    assert len(canonicals) == landing_count


def test_landing_pages_include_all_structured_data(client):
    for page in seo.landing_pages():
        body = client.get(page["path"]).get_data(as_text=True)
        assert {"HowTo", "FAQPage", "BreadcrumbList"} <= _jsonld_types(body), page["path"]


def test_landing_page_title_differs_from_home(client):
    home_title = _title(client.get("/").get_data(as_text=True))
    guitar_title = _title(client.get("/guitar-chord-transposer").get_data(as_text=True))
    assert guitar_title != home_title


def test_instrument_pages_exist(client):
    for slug in ("guitar-chord-transposer", "ukulele-chord-transposer", "piano-chord-transposer"):
        assert client.get(f"/{slug}").status_code == 200


def test_chromatic_chart_page_renders_grid(client):
    body = client.get("/chromatic-chart").get_data(as_text=True)
    assert "Chromatic transposition chart" in body
    assert "+7" in body


def test_chromatic_chart_page_shows_chart_before_transposer(client):
    body = client.get("/chromatic-chart").get_data(as_text=True)
    assert body.index("Every key, shifted by semitone") < body.index('id="panel-paste"')


def test_chromatic_chart_page_has_reference_tables(client):
    body = client.get("/chromatic-chart").get_data(as_text=True)
    assert "Semitones to interval names" in body
    assert "Perfect fifth" in body
    assert "Key signatures" in body
    assert "Relative minor" in body
    assert "F#m" in body


def test_key_pair_page_preselects_keys(client):
    body = client.get("/transpose/g-to-a").get_data(as_text=True)
    assert 'data-preselect-from="G"' in body
    assert 'data-preselect-to="A"' in body
    current_select = re.search(r'id="text_current_key".*?</select>', body, re.S).group(0)
    target_select = re.search(r'id="text_target_key".*?</select>', body, re.S).group(0)
    assert 'value="G" selected' in current_select
    assert 'value="A"  selected' in target_select


def test_uncurated_key_pair_returns_404(client):
    assert client.get("/transpose/c-to-c").status_code == 404
    assert client.get("/transpose/z-to-x").status_code == 404


def test_home_has_howto_jsonld(client):
    body = client.get("/").get_data(as_text=True)
    assert "HowTo" in _jsonld_types(body)
    assert "FAQPage" in _jsonld_types(body)


def test_faq_is_single_sourced(client):
    body = client.get("/").get_data(as_text=True)
    accordion_count = body.count("<summary")
    faq_block = next(block for block in _jsonld_blocks(body) if block.get("@type") == "FAQPage")
    assert accordion_count == len(faq_block["mainEntity"]) == len(seo.FAQ_ITEMS)


def test_home_has_no_capo_calculator(client):
    body = client.get("/").get_data(as_text=True)
    assert 'id="capo-calculator"' not in body
    assert 'id="capo-key"' not in body


def test_guitar_page_has_capo_calculator(client):
    body = client.get("/guitar-chord-transposer").get_data(as_text=True)
    assert 'id="capo-calculator"' in body
    assert 'id="capo-key"' in body


def test_instrument_pages_preselect_matching_instrument(client):
    expected = {
        "guitar-chord-transposer": "guitar",
        "ukulele-chord-transposer": "ukulele",
        "piano-chord-transposer": "piano",
    }
    for slug, instrument in expected.items():
        body = client.get(f"/{slug}").get_data(as_text=True)
        assert f'data-instrument="{instrument}"' in body


def test_home_and_key_pair_pages_omit_data_instrument(client):
    assert "data-instrument=" not in client.get("/").get_data(as_text=True)
    assert "data-instrument=" not in client.get("/transpose/g-to-a").get_data(as_text=True)


def test_file_page_exists_with_upload_panel(client):
    response = client.get("/file-chord-transposer")
    assert response.status_code == 200
    body = response.get_data(as_text=True)
    assert 'id="panel-upload"' in body
    assert 'id="file"' in body
    assert 'id="panel-paste"' not in body


def test_file_page_has_distinct_title_and_description(client):
    body = client.get("/file-chord-transposer").get_data(as_text=True)
    home_body = client.get("/").get_data(as_text=True)
    assert _title(body) != _title(home_body)
    description = re.search(r'<meta name="description" content="(.*?)"', body).group(1)
    assert ".docx" in description


def test_pages_have_page_specific_faq(client):
    cases = {
        "/guitar-chord-transposer": seo.GUITAR_FAQ_ITEMS,
        "/ukulele-chord-transposer": seo.UKULELE_FAQ_ITEMS,
        "/piano-chord-transposer": seo.PIANO_FAQ_ITEMS,
        "/file-chord-transposer": seo.FILE_FAQ_ITEMS,
        "/chromatic-chart": seo.CHART_FAQ_ITEMS,
    }
    for path, faq_items in cases.items():
        body = client.get(path).get_data(as_text=True)
        faq_block = next(block for block in _jsonld_blocks(body) if block.get("@type") == "FAQPage")
        rendered_questions = {item["question"] for item in faq_items}
        jsonld_questions = {entry["name"] for entry in faq_block["mainEntity"]}
        assert jsonld_questions == rendered_questions
        assert body.count("<summary") == len(faq_items)
        assert faq_items[0]["question"] in body


def test_nav_lists_file_transposer_first(client):
    sections = seo.nav_sections()
    assert sections[0]["links"][0]["path"] == "/file-chord-transposer"
    body = client.get("/").get_data(as_text=True)
    assert "/file-chord-transposer" in body


def test_nav_columns_balance_links_and_keep_all_sections():
    columns = seo.nav_columns()
    assert len(columns) == 2
    assert all(column for column in columns)
    flattened = [section for column in columns for section in column]
    assert sorted(s["heading"] for s in flattened) == sorted(
        s["heading"] for s in seo.nav_sections()
    )
    tall_column = max(columns, key=lambda column: sum(len(s["links"]) for s in column))
    assert any(s["heading"] == "By instrument" for s in tall_column)


def test_sitemap_lists_every_landing_url(client):
    body = client.get("/sitemap.xml").get_data(as_text=True)
    for path in seo.sitemap_paths():
        assert f"<loc>{seo.SITE_URL}{path}</loc>" in body
    assert body.count("<url>") == len(seo.sitemap_paths())


def test_robots_txt_points_to_sitemap(client):
    body = client.get("/robots.txt").get_data(as_text=True)
    assert f"Sitemap: {seo.SITE_URL}/sitemap.xml" in body


def test_capo_suggestion_covers_all_branches():
    assert seo._capo_suggestion("C", "C", 0) == "no capo is needed because the keys are the same"
    assert "fret 2" in seo._capo_suggestion("C", "D", 2)
    downshift = seo._capo_suggestion("C", "A", 9)
    assert "transpose down 3 semitones" in downshift


def _text_upload(data_bytes, filename="song.txt"):
    return {"file": (BytesIO(data_bytes), filename)}


def test_transpose_text_missing_file(client):
    response = client.post(
        "/transpose-text",
        data={"current_key": "C", "target_key": "D"},
        content_type="multipart/form-data",
    )
    assert response.status_code == 400
    assert "choose a .txt" in response.get_json()["error"]


def test_transpose_text_wrong_extension(client):
    data = _text_upload(b"C G", filename="song.docx")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Only .txt" in response.get_json()["error"]


def test_transpose_text_missing_keys(client):
    data = _text_upload(b"C G")
    data.update({"current_key": "", "target_key": ""})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "keys are required" in response.get_json()["error"]


def test_transpose_text_unsupported_format(client):
    data = _text_upload(b"C G")
    data.update({"current_key": "C", "target_key": "D", "format": "docx"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Unsupported output format" in response.get_json()["error"]


def test_transpose_text_empty_file(client):
    data = _text_upload(b"")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "empty" in response.get_json()["error"]


def test_transpose_text_too_large(client, monkeypatch):
    monkeypatch.setattr(app_module, "MAX_UPLOAD_BYTES", 4)
    data = _text_upload(b"C G Am F")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "too large" in response.get_json()["error"]


def test_transpose_text_invalid_encoding(client):
    data = _text_upload(b"\xff\xfe\x00chord")
    data.update({"current_key": "C", "target_key": "D"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "UTF-8" in response.get_json()["error"]


def test_transpose_text_invalid_key(client):
    data = _text_upload(b"C G")
    data.update({"current_key": "X", "target_key": "D"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "not a valid key" in response.get_json()["error"]


def test_transpose_text_txt_success(client):
    data = _text_upload(b"C G Am F")
    data.update({"current_key": "C", "target_key": "D", "format": "txt"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 200
    assert response.mimetype == "text/plain"
    assert "song_D.txt" in response.headers["Content-Disposition"]
    assert "C->D" in unquote(response.headers["X-Transpose-Changes"])


def test_transpose_text_same_key_reports_no_changes(client):
    data = _text_upload(b"C G Am F")
    data.update({"current_key": "C", "target_key": "C", "format": "txt"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 200
    assert unquote(response.headers["X-Transpose-Changes"]) == "none"


def test_transpose_text_chordpro_success(client):
    data = _text_upload(b"[C]Amazing [G]grace", filename="grace.cho")
    data.update({"current_key": "C", "target_key": "D", "format": "chordpro"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 200
    assert "grace_D.pro" in response.headers["Content-Disposition"]
    assert response.get_data(as_text=True).startswith("[D]Amazing")


def test_transpose_text_pdf_success(client, monkeypatch):
    monkeypatch.setattr(app_module, "convert_docx_to_pdf", lambda docx_bytes: _one_page_pdf())
    data = _text_upload(b"C G Am F")
    data.update({"current_key": "C", "target_key": "D", "format": "pdf"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 200
    assert response.mimetype == "application/pdf"
    assert "song_D.pdf" in response.headers["Content-Disposition"]


def test_transpose_text_pdf_conversion_failure(client, monkeypatch):
    def _raise(_docx_bytes):
        raise app_module.PdfConversionError("PDF conversion timed out.")

    monkeypatch.setattr(app_module, "convert_docx_to_pdf", _raise)
    data = _text_upload(b"C G Am F")
    data.update({"current_key": "C", "target_key": "D", "format": "pdf"})
    response = client.post("/transpose-text", data=data, content_type="multipart/form-data")
    assert response.status_code == 503
    assert "timed out" in response.get_json()["error"]


def _make_chord_docx(text="C G Am F"):
    document = Document()
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_binder_no_files(client):
    response = client.post(
        "/binder",
        data={"current_key": "C", "target_key": "D"},
        content_type="multipart/form-data",
    )
    assert response.status_code == 400
    assert "at least one" in response.get_json()["error"]


def test_binder_too_many_files(client, monkeypatch):
    monkeypatch.setattr(app_module, "MAX_BINDER_FILES", 1)
    data = {
        "files": [(BytesIO(_make_chord_docx()), "a.docx"), (BytesIO(_make_chord_docx()), "b.docx")],
        "current_key": "C",
        "target_key": "D",
    }
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "Too many files" in response.get_json()["error"]


def test_binder_missing_keys(client):
    data = {"files": [(BytesIO(_make_chord_docx()), "a.docx")], "current_key": "", "target_key": ""}
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "keys are required" in response.get_json()["error"]


def test_binder_empty_member(client):
    data = {
        "files": [(BytesIO(b""), "empty.docx")],
        "current_key": "C",
        "target_key": "D",
    }
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "empty" in response.get_json()["error"]


def test_binder_total_too_large(client, monkeypatch):
    monkeypatch.setattr(app_module, "MAX_BINDER_TOTAL_BYTES", 5)
    data = {
        "files": [(BytesIO(_make_chord_docx()), "a.docx")],
        "current_key": "C",
        "target_key": "D",
    }
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "combined maximum" in response.get_json()["error"]


def test_binder_invalid_key(client, monkeypatch):
    monkeypatch.setattr(app_module, "convert_docx_to_pdf", lambda docx_bytes: _one_page_pdf())
    data = {
        "files": [(BytesIO(_make_chord_docx()), "a.docx")],
        "current_key": "X",
        "target_key": "D",
    }
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 400
    assert "not a valid key" in response.get_json()["error"]


def test_binder_pdf_conversion_failure(client, monkeypatch):
    def _raise(_docx_bytes):
        raise app_module.PdfConversionError("Could not reach the PDF conversion service.")

    monkeypatch.setattr(app_module, "convert_docx_to_pdf", _raise)
    data = {
        "files": [(BytesIO(_make_chord_docx()), "a.docx")],
        "current_key": "C",
        "target_key": "D",
    }
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 503
    assert "conversion service" in response.get_json()["error"]


def test_binder_success_merges_pages(client, monkeypatch):
    monkeypatch.setattr(app_module, "convert_docx_to_pdf", lambda docx_bytes: _one_page_pdf())
    data = {
        "files": [
            (BytesIO(_make_chord_docx("C G")), "a.docx"),
            (BytesIO(_make_chord_docx("F C")), "b.docx"),
        ],
        "current_key": "C",
        "target_key": "D",
    }
    response = client.post("/binder", data=data, content_type="multipart/form-data")
    assert response.status_code == 200
    assert response.mimetype == "application/pdf"
    assert len(PdfReader(BytesIO(response.data)).pages) == 2
    assert "chord_binder_D.pdf" in response.headers["Content-Disposition"]
