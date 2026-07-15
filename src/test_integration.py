"""Integration tests for transpose.py using real .docx documents."""
# pylint: disable=missing-function-docstring,missing-class-docstring

from io import BytesIO

import pytest
from docx import Document

import transpose as t


def _make_docx(*paragraphs):
    """Build a real .docx in memory from a sequence of paragraph strings."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(*cell_texts):
    """Build a .docx whose body is a single-row table with the given cell texts."""
    doc = Document()
    table = doc.add_table(rows=1, cols=len(cell_texts))
    for cell, text in zip(table.rows[0].cells, cell_texts, strict=False):
        cell.paragraphs[0].add_run(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_header_footer(body_text, header_text, footer_text):
    """Build a .docx with body, header, and footer paragraphs."""
    doc = Document()
    doc.add_paragraph(body_text)
    section = doc.sections[0]
    section.header.paragraphs[0].add_run(header_text)
    section.footer.paragraphs[0].add_run(footer_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _paragraphs(docx_bytes):
    """Return the text of every non-empty paragraph in a .docx bytes object."""
    return [p.text for p in Document(BytesIO(docx_bytes)).paragraphs if p.text]


def _chord_semitones_in_line(text, german=False):
    """Return the set of semitone values for every root note in a chord line."""
    semitones = set()
    for token in text.split():
        core = token.strip("().,;:…\"'-")
        m = t.ROOT_MATCH_RE.match(core)
        if m:
            try:
                semitones.add(t.note_semitone(m.group(1), german))
            except KeyError:
                pass
    return semitones


# ---------------------------------------------------------------------------
# German detection edge cases
# ---------------------------------------------------------------------------


class TestGermanDetection:
    def test_h_chord_triggers_german(self):
        file_bytes = _make_docx("C G H F")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True

    def test_es_chord_without_h_triggers_german(self):
        """A sheet using Es (Eb) but no H is still detected as German."""
        file_bytes = _make_docx("C G Es F")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True

    def test_as_chord_without_h_triggers_german(self):
        file_bytes = _make_docx("C G As F")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True

    def test_cis_chord_without_h_triggers_german(self):
        file_bytes = _make_docx("Cis Fis Gism")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True

    def test_fis_chord_without_h_triggers_german(self):
        file_bytes = _make_docx("D A Fis Gm")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True

    def test_mixed_h_and_sharp_notation_is_german(self):
        """Sheet using both H and C# is detected as German."""
        file_bytes = _make_docx("C# G#m H F#")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True

    def test_english_sharp_sheet_without_german_markers_not_detected(self):
        """A sheet with C#, G#m, B, F# but no H/Es/As/Cis is NOT detected as German.

        This is a known limitation: the app cannot distinguish a German sheet that
        uses only '#' accidentals from an English sheet. In this case 'B' is treated
        as B natural (semitone 11), not B-flat (semitone 10).
        """
        file_bytes = _make_docx("C# G#m B F#")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is False

    def test_b_in_non_german_doc_treated_as_b_natural(self):
        """When detect_german returns False, B transposes as B natural (11), not Bb (10)."""
        # Transpose C# -> E (3 semitones). B natural (11) + 3 = 14 % 12 = 2 = D.
        # If it were B-flat (10) + 3 = 13 % 12 = 1 = C#/Db.
        file_bytes = _make_docx("C# B")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C#", "E")
        assert ("B", "D") in changes  # B treated as B natural

    def test_lyric_line_containing_h_word_does_not_trigger_german(self):
        """'Hello' on a lyrics line must not be mis-detected as the German H chord."""
        file_bytes = _make_docx("Hello world, He said")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is False

    def test_german_detected_from_second_paragraph(self):
        """detect_german must scan all paragraphs, not just the first."""
        file_bytes = _make_docx("C G Am F", "Es B H")
        doc = Document(BytesIO(file_bytes))
        assert t.detect_german(doc) is True


# ---------------------------------------------------------------------------
# German sheet: B = Bb, H = B natural
# ---------------------------------------------------------------------------


class TestGermanBAndH:
    def test_b_in_german_doc_transposes_as_bb(self):
        """In a German document, B (Bb) + 2 semitones = C (Bb->B natural->C)."""
        # Sheet has H so it's detected as German. B = Bb (10), +2 = 0 = C.
        file_bytes = _make_docx("H B C")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert ("B", "C") in changes  # B(Bb=10) + 2 = 12%12=0 -> C
        assert ("H", "Cis") in changes  # H(B=11) + 2 = 13%12=1 -> Cis (sharp spelling for D)

    def test_h_in_german_doc_transposes_as_b_natural(self):
        """H is B natural (11). H + 1 semitone = C."""
        file_bytes = _make_docx("H G")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "H", "C")
        assert ("H", "C") in changes

    def test_german_b_not_confused_with_english_b(self):
        """Transposing a German doc: B (=Bb) and H (=B natural) produce different results."""
        file_bytes = _make_docx("H B")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        original_chords = {orig for orig, _ in changes}
        assert "H" in original_chords
        assert "B" in original_chords
        h_result = next(new for orig, new in changes if orig == "H")
        b_result = next(new for orig, new in changes if orig == "B")
        assert h_result != b_result  # They should land on different semitones


# ---------------------------------------------------------------------------
# Round-trip transposition
# ---------------------------------------------------------------------------


class TestRoundTrip:
    def _round_trip(self, chord_line, from_key, to_key):
        file_bytes = _make_docx(chord_line)
        mid_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, from_key, to_key)
        result_bytes, _, _, _ = t.transpose_document_bytes(mid_bytes, to_key, from_key)
        return _paragraphs(result_bytes)[0]

    def test_c_to_g_and_back(self):
        assert self._round_trip("C G Am F", "C", "G") == "C G Am F"

    def test_c_to_gb_and_back(self):
        """Round-tripping through Gb (tritone) restores the original chord line.

        Extra spaces between chords ensure that chord-length growth never collapses
        adjacent tokens into one unrecognised token, which would break the second pass.
        """
        file_bytes = _make_docx("C    G    Am   F")
        mid_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "Gb")
        result_bytes, _, _, _ = t.transpose_document_bytes(mid_bytes, "Gb", "C")
        assert _paragraphs(result_bytes)[0] == "C    G    Am   F"

    def test_all_12_semitones_round_trip(self):
        """Twelve consecutive half-step-up transpositions return to the original chord line.

        Wide spacing prevents chord-length growth from collapsing adjacent tokens.
        We step through each chromatic key so current_key always matches the document.
        """
        chromatic = ["C", "C#", "D", "Eb", "E", "F", "F#", "G", "Ab", "A", "Bb", "B"]
        original_line = "C    Dm   Em   F    G    Am"
        file_bytes = _make_docx(original_line)
        current = file_bytes
        for i in range(12):
            from_key = chromatic[i % 12]
            to_key = chromatic[(i + 1) % 12]
            current, _, _, _ = t.transpose_document_bytes(current, from_key, to_key)
        assert _paragraphs(current)[0].rstrip() == original_line

    def test_minor_key_round_trip(self):
        assert self._round_trip("Am Dm E7 Am", "Am", "Em") == "Am Dm E7 Am"

    def test_flat_key_round_trip(self):
        """Round-tripping a flat-key sheet restores the original chord line."""
        file_bytes = _make_docx("Bb   Eb   F    Gm")
        mid_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "Bb", "Eb")
        result_bytes, _, _, _ = t.transpose_document_bytes(mid_bytes, "Eb", "Bb")
        assert _paragraphs(result_bytes)[0] == "Bb   Eb   F    Gm"

    def test_german_round_trip(self):
        """A German sheet round-trips correctly through transposition."""
        file_bytes = _make_docx("H Es B Cis")
        mid_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "H", "C")
        result_bytes, _, _, _ = t.transpose_document_bytes(mid_bytes, "C", "H")
        para = _paragraphs(result_bytes)[0]
        assert "H" in para
        assert "Es" in para
        assert "B" in para


# ---------------------------------------------------------------------------
# Enharmonic equivalence
# ---------------------------------------------------------------------------


class TestEnharmonicEquivalence:
    def test_c_sharp_and_db_source_keys_give_same_semitone_distance(self):
        """Transposing from C# and from Db to the same target give the same pitch classes."""
        file_sharp = _make_docx("C# F# G#m")
        file_flat = _make_docx("Db Gb Abm")
        result_sharp, _, _, _ = t.transpose_document_bytes(file_sharp, "C#", "E")
        result_flat, _, _, _ = t.transpose_document_bytes(file_flat, "Db", "E")
        semitones_sharp = _chord_semitones_in_line(_paragraphs(result_sharp)[0])
        semitones_flat = _chord_semitones_in_line(_paragraphs(result_flat)[0])
        assert semitones_sharp == semitones_flat

    def test_f_sharp_and_gb_as_target_keys_use_correct_spelling(self):
        file_bytes = _make_docx("C G Am F")
        _, _, to_sharp, changes_sharp = t.transpose_document_bytes(file_bytes, "C", "F#")
        _, _, to_flat, changes_flat = t.transpose_document_bytes(file_bytes, "C", "Gb")
        assert to_sharp == "F#"
        assert to_flat == "Gb"
        sharp_notes = {new for _, new in changes_sharp}
        flat_notes = {new for _, new in changes_flat}
        assert any("#" in n for n in sharp_notes)
        assert not any("#" in n for n in flat_notes)

    def test_bb_and_a_sharp_as_target_give_same_pitches(self):
        """Bb and A# are the same pitch; the resulting chord pitch classes must match."""
        file_bytes = _make_docx("C G Am F")
        result_bb, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "Bb")
        result_as, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "A#")
        semitones_bb = _chord_semitones_in_line(_paragraphs(result_bb)[0])
        semitones_as = _chord_semitones_in_line(_paragraphs(result_as)[0])
        assert semitones_bb == semitones_as


# ---------------------------------------------------------------------------
# Chord quality preservation
# ---------------------------------------------------------------------------


class TestChordQualitiesPreserved:
    def _transpose_and_get_para(self, chord_line, from_key="C", to_key="D"):
        file_bytes = _make_docx(chord_line)
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, from_key, to_key)
        return _paragraphs(result_bytes)[0]

    def test_maj7_quality_preserved(self):
        result = self._transpose_and_get_para("Cmaj7 Fmaj7")
        assert "maj7" in result

    def test_dim7_quality_preserved(self):
        result = self._transpose_and_get_para("Bdim7")
        assert "dim7" in result

    def test_aug_quality_preserved(self):
        result = self._transpose_and_get_para("Caug")
        assert "aug" in result

    def test_sus4_quality_preserved(self):
        result = self._transpose_and_get_para("Gsus4")
        assert "sus4" in result

    def test_add9_quality_preserved(self):
        result = self._transpose_and_get_para("Cadd9")
        assert "add9" in result

    def test_m7_quality_preserved(self):
        result = self._transpose_and_get_para("Am7 Dm7")
        assert "m7" in result

    def test_slash_chord_quality_preserved(self):
        result = self._transpose_and_get_para("Cmaj7/E")
        assert "maj7" in result
        assert "/" in result

    def test_slash_chord_both_notes_transposed(self):
        """Both root and bass note of a slash chord are transposed by the same interval."""
        result = self._transpose_and_get_para("G/B")  # C->D: G->A, B->C#
        assert result.strip() == "A/C#"

    def test_complex_chord_line(self):
        result = self._transpose_and_get_para("Cmaj7 Am7 Dm7 G7sus4")
        assert "maj7" in result
        assert "m7" in result
        assert "sus4" in result


# ---------------------------------------------------------------------------
# Lyrics lines never change
# ---------------------------------------------------------------------------


class TestLyricsPreserved:
    def test_pure_lyric_line_unchanged(self):
        file_bytes = _make_docx("C G Am F", "These are the lyrics")
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "D")
        paras = _paragraphs(result_bytes)
        assert "These are the lyrics" in paras

    def test_only_chord_lines_produce_changes(self):
        file_bytes = _make_docx("Just some words here", "Another lyric line")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert changes == []

    def test_mixed_document_chord_lines_transposed_lyrics_not(self):
        file_bytes = _make_docx(
            "C G Am F",
            "Amazing grace how sweet the sound",
            "F C G",
            "That saved a wretch like me",
        )
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "D")
        paras = _paragraphs(result_bytes)
        assert "Amazing grace how sweet the sound" in paras
        assert "That saved a wretch like me" in paras
        assert "C G Am F" not in paras

    def test_section_label_line_with_chords_is_transposed(self):
        """Lines like 'REF: C G Am' are chord lines and should be transposed."""
        file_bytes = _make_docx("REF: C G Am")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert any(orig == "C" for orig, _ in changes)

    def test_line_with_stray_chord_word_among_lyrics_not_transposed(self):
        """'C the G' is not a chord line — the word 'the' disqualifies it."""
        file_bytes = _make_docx("C the G")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert changes == []


# ---------------------------------------------------------------------------
# Document structure: tables and headers/footers
# ---------------------------------------------------------------------------


class TestDocumentStructure:
    def test_chords_in_table_cells_are_transposed(self):
        file_bytes = _make_docx_with_table("C G Am F", "Verse lyrics")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert any(orig == "C" for orig, _ in changes)

    def test_chords_in_header_are_transposed(self):
        file_bytes = _make_docx_with_header_footer("Lyrics line", "C G Am F", "Footer text")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert any(orig == "C" for orig, _ in changes)

    def test_chords_in_footer_are_transposed(self):
        file_bytes = _make_docx_with_header_footer("Lyrics line", "Header text", "C G Am F")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert any(orig == "C" for orig, _ in changes)

    def test_empty_document_produces_no_changes(self):
        file_bytes = _make_docx("")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert changes == []

    def test_document_with_only_empty_paragraphs(self):
        file_bytes = _make_docx("", "", "")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert changes == []


# ---------------------------------------------------------------------------
# German document spelling (sharp vs. flat output)
# ---------------------------------------------------------------------------


class TestGermanSpelling:
    def test_german_doc_to_flat_key_uses_german_flat_names(self):
        """Transposing a German sheet to a flat key should produce Es, As, Des etc."""
        file_bytes = _make_docx("H C D")  # H triggers German detection
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "Es")
        para = _paragraphs(result_bytes)[0]
        # In the result, flat accidentals should be German-style (Es, As, Des)
        assert "b" not in para  # no English-style flats

    def test_german_doc_to_sharp_key_uses_german_sharp_names(self):
        """Transposing a German sheet to a sharp key should produce Cis, Fis, Gis etc."""
        file_bytes = _make_docx("C Es As H")  # Es/As trigger German detection
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "D")
        para = _paragraphs(result_bytes)[0]
        # D is a sharp key; any accidentals should be 'is' style, not 'b'
        assert "b" not in para

    def test_english_doc_to_flat_key_uses_english_flat_names(self):
        file_bytes = _make_docx("C G Am F")
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "Eb")
        para = _paragraphs(result_bytes)[0]
        assert "Eb" in para or "Bb" in para or "Ab" in para

    def test_english_doc_to_sharp_key_uses_english_sharp_names(self):
        file_bytes = _make_docx("C G Am F")
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C", "E")
        para = _paragraphs(result_bytes)[0]
        assert "b" not in para  # E major uses sharps


# ---------------------------------------------------------------------------
# Same-key transposition (no-op)
# ---------------------------------------------------------------------------


class TestSameKeyNoOp:
    def test_same_key_changes_empty(self):
        file_bytes = _make_docx("C G Am F")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "C")
        assert changes == []

    def test_same_key_document_content_unchanged(self):
        file_bytes = _make_docx("C G Am F", "Lyrics here")
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "G", "G")
        paras = _paragraphs(result_bytes)
        assert "C G Am F" in paras
        assert "Lyrics here" in paras

    def test_enharmonic_same_pitch_respells_but_does_not_shift(self):
        """C# and Db are the same pitch class (0 semitone shift).

        The app re-spells chord names to match the target key's convention, so
        chords will be renamed (C# -> Db) even though no pitch shift occurs.
        The resulting pitch classes must equal the originals.
        """
        file_bytes = _make_docx("C# F# G#m")
        result_bytes, _, _, _ = t.transpose_document_bytes(file_bytes, "C#", "Db")
        original_semitones = _chord_semitones_in_line(_paragraphs(file_bytes)[0])
        result_semitones = _chord_semitones_in_line(_paragraphs(result_bytes)[0])
        assert result_semitones == original_semitones


# ---------------------------------------------------------------------------
# Specific interval correctness
# ---------------------------------------------------------------------------


class TestIntervalCorrectness:
    def test_up_one_semitone_c_to_db(self):
        file_bytes = _make_docx("C F G")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "Db")
        assert ("C", "Db") in changes
        assert ("F", "Gb") in changes
        assert ("G", "Ab") in changes

    def test_up_perfect_fifth_c_to_g(self):
        file_bytes = _make_docx("C Am F G")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "G")
        assert ("C", "G") in changes
        assert ("Am", "Em") in changes
        assert ("F", "C") in changes
        assert ("G", "D") in changes

    def test_down_perfect_fifth_equals_up_seven(self):
        """C down a fifth = F, same as C up 7 semitones."""
        file_bytes = _make_docx("C G Am F")
        result_down, _, _, _ = t.transpose_document_bytes(file_bytes, "G", "C")
        paras = _paragraphs(result_down)
        assert "C" in paras[0]

    def test_tritone_transposition_c_to_gb(self):
        file_bytes = _make_docx("C Dm G")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "Gb")
        assert ("C", "Gb") in changes

    def test_octave_is_identity(self):
        """12 semitones = 1 octave = no change in pitch class."""
        file_bytes = _make_docx("C G Am F")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "C")
        assert changes == []


# ---------------------------------------------------------------------------
# Metadata returned by transpose_document_bytes
# ---------------------------------------------------------------------------


class TestMetadata:
    def test_from_label_is_correct(self):
        file_bytes = _make_docx("Am Dm E")
        _, from_label, _, _ = t.transpose_document_bytes(file_bytes, "Am", "Em")
        assert from_label == "Am"

    def test_to_label_is_correct(self):
        file_bytes = _make_docx("Am Dm E")
        _, _, to_label, _ = t.transpose_document_bytes(file_bytes, "Am", "Em")
        assert to_label == "Em"

    def test_changes_are_sorted(self):
        file_bytes = _make_docx("C G Am F")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert changes == sorted(changes)

    def test_changes_contain_no_duplicates(self):
        file_bytes = _make_docx("C G Am F", "C G Am F")
        _, _, _, changes = t.transpose_document_bytes(file_bytes, "C", "D")
        assert len(changes) == len(set(changes))

    def test_invalid_current_key_raises(self):
        file_bytes = _make_docx("C G")
        with pytest.raises(t.InvalidKeyError, match="not a valid key"):
            t.transpose_document_bytes(file_bytes, "X", "D")

    def test_invalid_target_key_raises(self):
        file_bytes = _make_docx("C G")
        with pytest.raises(t.InvalidKeyError, match="not a valid key"):
            t.transpose_document_bytes(file_bytes, "C", "X")

    def test_german_key_label_in_from_label(self):
        file_bytes = _make_docx("H Es B")
        _, from_label, _, _ = t.transpose_document_bytes(file_bytes, "H", "C")
        assert from_label == "H"

    def test_minor_labels_carry_m_suffix(self):
        file_bytes = _make_docx("Am Em F C")
        _, from_label, to_label, _ = t.transpose_document_bytes(file_bytes, "Am", "Bm")
        assert from_label == "Am"
        assert to_label == "Bm"


class TestTranposeTextEndToEnd:
    def test_multiline_sheet_matches_docx_path(self):
        source = "C       G\nAm      F\nLyrics line here\n\nF       C"
        text_result, _, _, text_changes = t.transpose_text(source, "C", "D")

        docx_bytes = _make_docx(*source.split("\n"))
        docx_out, _, _, docx_changes = t.transpose_document_bytes(docx_bytes, "C", "D")
        docx_lines = _paragraphs(docx_out)

        transposed_chord_lines = [line for line in text_result.split("\n") if line.strip()]
        assert docx_lines == transposed_chord_lines
        assert text_changes == docx_changes
        assert "Lyrics line here" in text_result

    def test_german_sheet_stays_german(self):
        source = "H Es B\nFis Cis Gis"
        result, from_label, to_label, _ = t.transpose_text(source, "H", "C")
        assert from_label == "H"
        assert to_label == "C"
        assert result.split("\n")[0].startswith("C")

    def test_spacing_preserved_across_chord_width_change(self):
        source = "C   F#  G"
        result, _, _, _ = t.transpose_text(source, "C", "Db")
        assert result.split("\n")[0].startswith("Db")


class TestChordProRoundTrip:
    def test_chordpro_round_trips(self):
        source = "[C]Amazing [F]grace how [G]sweet the [C]sound"
        mid, _, _, _ = t.transpose_chordpro_text(source, "C", "G")
        result, _, _, _ = t.transpose_chordpro_text(mid, "G", "C")
        assert result == source

    def test_chordpro_lyrics_preserved(self):
        source = "[C]Amazing [G]grace\nplain lyric line\n[F]how sweet"
        result, _, _, _ = t.transpose_chordpro_text(source, "C", "D")
        assert "plain lyric line" in result
        assert result.startswith("[D]Amazing [A]grace")

    def test_chordpro_to_nashville_and_plain(self):
        source = "[C]Amazing [G]grace"
        nashville, tonic = t.text_to_nashville(source, "C")
        assert nashville == "[1]Amazing [5]grace"
        assert tonic == "C"
        assert t._chordpro_to_plain(source) == "C G"
